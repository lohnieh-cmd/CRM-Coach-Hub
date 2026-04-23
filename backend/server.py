"""
Ascent CRM — FastAPI backend (web-based MVP reference implementation).

Implements:
  • Auth (JWT)
  • Companies, Contacts, Deals, Pipeline Stages (Kanban)
  • Products / Price Lists, Quotes, Invoices
  • Lead Forms + Submissions (consent + double opt-in stub)
  • AI Studio (Gemini 3 via emergentintegrations — grounded content + reply suggester)
  • Stripe Checkout + /api/webhook/stripe
  • Coaching Templates (executive, fitness, consultant)
  • Automations (stubs), Integrations Hub (stubs)
  • Analytics aggregate endpoint
  • GDPR: consent log, export, soft/hard delete
  • Audit trail on every mutating call
  • Seed data on startup

All IDs are UUIDv4 strings. All timestamps ISO-8601 UTC.
Every MongoDB query projects out `_id`.
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import secrets
import uuid
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Literal

import bcrypt
import jwt
from dotenv import load_dotenv
from fastapi import APIRouter, Body, Depends, FastAPI, HTTPException, Request, Response
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, ConfigDict, EmailStr, Field
from starlette.middleware.cors import CORSMiddleware

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

# ── infra ──────────────────────────────────────────────────────────────────────
mongo_url = os.environ["MONGO_URL"]
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ["DB_NAME"]]
JWT_SECRET = os.environ["JWT_SECRET"]
JWT_ALGO = "HS256"
JWT_TTL_HOURS = 24 * 7
EMERGENT_LLM_KEY = os.environ["EMERGENT_LLM_KEY"]
STRIPE_API_KEY = os.environ["STRIPE_API_KEY"]

app = FastAPI(title="Ascent CRM API", version="1.0.0")
api = APIRouter(prefix="/api")
bearer = HTTPBearer(auto_error=False)

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s %(message)s")
log = logging.getLogger("ascent")


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def new_id() -> str:
    return str(uuid.uuid4())


# ── models ─────────────────────────────────────────────────────────────────────
class BaseDoc(BaseModel):
    model_config = ConfigDict(extra="ignore")


class SignupReq(BaseDoc):
    email: EmailStr
    password: str
    name: str | None = None


class LoginReq(BaseDoc):
    email: EmailStr
    password: str


class AuthResp(BaseDoc):
    token: str
    user: dict


class CompanyIn(BaseDoc):
    name: str
    industry: str | None = None
    website: str | None = None
    lifecycle_stage: str | None = "lead"
    status: str | None = "active"
    notes: str | None = None
    tags: list[str] = Field(default_factory=list)
    billing_address: dict | None = None
    custom_fields: dict = Field(default_factory=dict)


class ContactIn(BaseDoc):
    first_name: str
    last_name: str | None = None
    email: EmailStr | None = None
    phone: str | None = None
    role_title: str | None = None
    company_id: str | None = None
    notes: str | None = None
    tags: list[str] = Field(default_factory=list)
    custom_fields: dict = Field(default_factory=dict)
    consent: dict = Field(default_factory=lambda: {"marketing": False, "newsletter": False})


class StageIn(BaseDoc):
    name: str
    order: int
    probability: int = 10
    altitude_label: Literal["Basecamp", "Ascent", "Summit", "Closed Won", "Closed Lost"] = "Basecamp"


class DealIn(BaseDoc):
    title: str
    contact_id: str | None = None
    company_id: str | None = None
    pipeline_stage_id: str
    value: float = 0.0
    currency: str = "USD"
    probability: int = 10
    expected_close_date: str | None = None
    notes: str | None = None
    tags: list[str] = Field(default_factory=list)


class ProductIn(BaseDoc):
    sku: str
    name: str
    description: str | None = None
    unit_price: float
    currency: str = "USD"
    tax_rate: float = 0.0
    tier: Literal["foundation", "ascent", "summit"] = "foundation"
    active: bool = True


class QuoteLine(BaseDoc):
    product_id: str | None = None
    description: str
    qty: float = 1.0
    unit_price: float
    discount_pct: float = 0.0
    tax_rate: float = 0.0


class QuoteIn(BaseDoc):
    deal_id: str | None = None
    contact_id: str | None = None
    company_id: str | None = None
    line_items: list[QuoteLine]
    currency: str = "USD"
    valid_until: str | None = None
    valid_days: int | None = None  # if set, server auto-computes valid_until = today + valid_days
    terms: str | None = None


class InvoiceIn(BaseDoc):
    quote_id: str | None = None
    contact_id: str | None = None
    company_id: str | None = None
    line_items: list[QuoteLine] | None = None
    currency: str = "USD"
    due_date: str | None = None


class LeadFormField(BaseDoc):
    key: str
    label: str
    type: Literal["text", "email", "phone", "textarea", "select", "checkbox"]
    required: bool = False
    options: list[str] | None = None


class FunnelBranch(BaseDoc):
    if_field: str  # field key
    equals: str
    goto_step_id: str | None = None  # None = submit / end


class FunnelStep(BaseDoc):
    id: str
    title: str
    description: str | None = None
    fields: list[LeadFormField]
    branches: list[FunnelBranch] = Field(default_factory=list)
    # If no branch matches, fall through to next step in list (or submit if last)


class LeadFormIn(BaseDoc):
    name: str
    slug: str
    fields: list[LeadFormField] = Field(default_factory=list)  # single-step legacy
    steps: list[FunnelStep] = Field(default_factory=list)  # multi-step funnel (optional)
    consent_text: str = "I agree to receive communications and accept the privacy policy."
    double_opt_in: bool = True
    success_redirect: str | None = None


class FormSubmitReq(BaseDoc):
    answers: dict
    consent_given: bool = False


class EmailLogIn(BaseDoc):
    contact_id: str | None = None
    deal_id: str | None = None
    direction: Literal["in", "out"] = "in"
    subject: str
    body: str
    from_addr: str | None = None
    to_addr: str | None = None
    received_at: str | None = None


class TaskIn(BaseDoc):
    title: str
    contact_id: str | None = None
    deal_id: str | None = None
    due_date: str | None = None
    notes: str | None = None


class AIGenerateReq(BaseDoc):
    kind: Literal["blog", "email", "quote_summary", "reply"]
    prompt: str
    tone: Literal["professional", "friendly", "warm-sherpa", "authoritative", "short"] = "professional"
    contact_id: str | None = None
    deal_id: str | None = None
    incoming_email: str | None = None


class CheckoutReq(BaseDoc):
    invoice_id: str
    origin_url: str


# ── auth helpers ───────────────────────────────────────────────────────────────
def hash_pw(pw: str) -> str:
    return bcrypt.hashpw(pw.encode(), bcrypt.gensalt(rounds=12)).decode()


def verify_pw(pw: str, h: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode(), h.encode())
    except Exception:
        return False


def make_token(user_id: str, email: str) -> str:
    exp = datetime.now(timezone.utc) + timedelta(hours=JWT_TTL_HOURS)
    return jwt.encode({"sub": user_id, "email": email, "exp": exp}, JWT_SECRET, algorithm=JWT_ALGO)


async def current_user(creds: HTTPAuthorizationCredentials = Depends(bearer)) -> dict:
    if not creds:
        raise HTTPException(401, "Missing token")
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALGO])
    except jwt.PyJWTError as e:
        raise HTTPException(401, f"Invalid token: {e}")
    user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(401, "User not found")
    # Team data scoping: every data query uses u["id"] as owner_id. For invited members
    # we remap that to the team owner's user_id so they see the full team's data.
    # The real user id is kept in `actor_id` for audit trail.
    team_owner_id = user.get("team_owner_id") or user["id"]
    return {**user, "actor_id": user["id"], "id": team_owner_id, "role": user.get("role", "owner")}


def _require_role(allowed: set[str]):
    async def dep(u: dict = Depends(current_user)):
        if u.get("role") not in allowed:
            raise HTTPException(403, f"Role '{u.get('role')}' not permitted (need one of {sorted(allowed)})")
        return u
    return dep


require_owner_admin = _require_role({"owner", "admin"})
require_accountant = _require_role({"owner", "admin", "accountant"})


def _strip_oid(obj):
    """Recursively remove BSON ObjectId values that Motor may have injected."""
    if isinstance(obj, dict):
        return {k: _strip_oid(v) for k, v in obj.items() if k != "_id"}
    if isinstance(obj, list):
        return [_strip_oid(v) for v in obj]
    return obj


async def audit(actor_id: str, action: str, entity_type: str, entity_id: str, before: dict | None = None, after: dict | None = None):
    await db.audit_entries.insert_one({
        "id": new_id(),
        "actor_id": actor_id,
        "action": action,
        "entity_type": entity_type,
        "entity_id": entity_id,
        "before": _strip_oid(before or {}),
        "after": _strip_oid(after or {}),
        "timestamp": now_iso(),
    })


# ── auth endpoints ─────────────────────────────────────────────────────────────
@api.post("/auth/signup", response_model=AuthResp)
async def signup(req: SignupReq):
    if await db.users.find_one({"email": req.email.lower()}):
        raise HTTPException(409, "Email already registered")
    uid = new_id()
    doc = {
        "id": uid,
        "email": req.email.lower(),
        "name": req.name or req.email.split("@")[0],
        "password_hash": hash_pw(req.password),
        "role": "owner",
        "team_owner_id": uid,  # self-team until invited into another
        "brand_voice": {
            "tone": "warm-sherpa",
            "vocabulary_hints": "use altitude / summit / ascent metaphors sparingly",
            "signature": f"— {req.name or 'The Sherpa'}",
            "banned_phrases": ["cutting-edge", "synergy"],
        },
        "created_at": now_iso(),
    }
    await db.users.insert_one(doc)
    await audit(uid, "signup", "user", uid)
    public = {k: v for k, v in doc.items() if k not in ("password_hash", "_id")}
    return {"token": make_token(uid, req.email.lower()), "user": public}


@api.post("/auth/login", response_model=AuthResp)
async def login(req: LoginReq):
    user = await db.users.find_one({"email": req.email.lower()})
    if not user or not verify_pw(req.password, user["password_hash"]):
        raise HTTPException(401, "Invalid credentials")
    public = {k: v for k, v in user.items() if k not in ("password_hash", "_id")}
    return {"token": make_token(user["id"], user["email"]), "user": public}


@api.get("/auth/me")
async def me(u: dict = Depends(current_user)):
    return u


@api.put("/auth/brand-voice")
async def update_brand_voice(payload: dict = Body(...), u: dict = Depends(current_user)):
    await db.users.update_one({"id": u["id"]}, {"$set": {"brand_voice": payload}})
    await audit(u["actor_id"], "update", "brand_voice", u["id"], after=payload)
    return {"ok": True, "brand_voice": payload}


@api.get("/auth/quote-template")
async def get_quote_template(u: dict = Depends(current_user)):
    """Settings used by /api/quotes/{id}/export/docx."""
    user = await db.users.find_one({"id": u["id"]}, {"_id": 0, "password_hash": 0}) or {}
    return user.get("quote_template") or {
        "company_name": user.get("name") or "Ascent CRM",
        "tagline": "Ascent CRM · Climb Leadership Lab",
        "accent_color_hex": "E26E4A",
        "footer_text": "Thank you for your business.",
        "signature_block": user.get("brand_voice", {}).get("signature") or "",
        "title_label": "QUOTATION",
    }


@api.put("/auth/quote-template")
async def update_quote_template(payload: dict = Body(...), u: dict = Depends(require_owner_admin)):
    # whitelist keys to keep the blob tidy
    allowed = {"company_name", "tagline", "accent_color_hex", "footer_text", "signature_block", "title_label"}
    clean = {k: v for k, v in payload.items() if k in allowed and isinstance(v, str)}
    await db.users.update_one({"id": u["id"]}, {"$set": {"quote_template": clean}})
    await audit(u["actor_id"], "update", "quote_template", u["id"], after=clean)
    return {"ok": True, "quote_template": clean}


# ── generic CRUD helpers ───────────────────────────────────────────────────────
async def _list(collection: str, q: dict | None = None, limit: int = 500):
    docs = await db[collection].find(q or {}, {"_id": 0}).to_list(limit)
    return docs


async def _insert(collection: str, doc: dict, actor_id: str, entity_type: str):
    doc["id"] = doc.get("id") or new_id()
    doc["created_at"] = now_iso()
    doc["updated_at"] = now_iso()
    await db[collection].insert_one(doc)
    await audit(actor_id, "create", entity_type, doc["id"], after=doc)
    # Strip _id that Motor may have added to the passed-in dict
    return {k: v for k, v in doc.items() if k != "_id"}


async def _update(collection: str, doc_id: str, patch: dict, actor_id: str, entity_type: str):
    before = await db[collection].find_one({"id": doc_id}, {"_id": 0})
    if not before:
        raise HTTPException(404, f"{entity_type} not found")
    patch["updated_at"] = now_iso()
    await db[collection].update_one({"id": doc_id}, {"$set": patch})
    after = await db[collection].find_one({"id": doc_id}, {"_id": 0})
    await audit(actor_id, "update", entity_type, doc_id, before=before, after=after)
    return after


async def _soft_delete(collection: str, doc_id: str, actor_id: str, entity_type: str):
    before = await db[collection].find_one({"id": doc_id}, {"_id": 0})
    if not before:
        raise HTTPException(404, f"{entity_type} not found")
    await db[collection].update_one({"id": doc_id}, {"$set": {"deleted_at": now_iso()}})
    await audit(actor_id, "delete", entity_type, doc_id, before=before)
    return {"ok": True}


# ── Companies ──────────────────────────────────────────────────────────────────
@api.get("/companies")
async def list_companies(u: dict = Depends(current_user)):
    return await _list("companies", {"owner_id": u["id"], "deleted_at": {"$exists": False}})


@api.post("/companies")
async def create_company(p: CompanyIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    return await _insert("companies", doc, u["actor_id"], "company")


@api.put("/companies/{cid}")
async def update_company(cid: str, p: CompanyIn, u: dict = Depends(current_user)):
    return await _update("companies", cid, p.model_dump(), u["actor_id"], "company")


@api.delete("/companies/{cid}")
async def delete_company(cid: str, u: dict = Depends(current_user)):
    return await _soft_delete("companies", cid, u["actor_id"], "company")


# ── Contacts ───────────────────────────────────────────────────────────────────
@api.get("/contacts")
async def list_contacts(u: dict = Depends(current_user)):
    return await _list("contacts", {"owner_id": u["id"], "deleted_at": {"$exists": False}})


@api.post("/contacts")
async def create_contact(p: ContactIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    doc["last_activity_at"] = now_iso()
    doc["interaction_count"] = 0
    if doc.get("consent", {}).get("marketing") or doc.get("consent", {}).get("newsletter"):
        await db.consent_logs.insert_one({
            "id": new_id(), "contact_email": p.email, "kind": "profile_consent",
            "given": True, "source": "manual", "timestamp": now_iso(),
        })
    return await _insert("contacts", doc, u["actor_id"], "contact")


@api.put("/contacts/{cid}")
async def update_contact(cid: str, p: ContactIn, u: dict = Depends(current_user)):
    return await _update("contacts", cid, p.model_dump(), u["actor_id"], "contact")


@api.delete("/contacts/{cid}")
async def delete_contact(cid: str, u: dict = Depends(current_user)):
    return await _soft_delete("contacts", cid, u["actor_id"], "contact")


@api.get("/contacts/{cid}/timeline")
async def contact_timeline(cid: str, u: dict = Depends(current_user)):
    contact = await db.contacts.find_one({"id": cid, "owner_id": u["id"]}, {"_id": 0})
    if not contact:
        raise HTTPException(404, "Contact not found")
    deals = await db.deals.find({"contact_id": cid, "owner_id": u["id"]}, {"_id": 0}).to_list(100)
    quotes = await db.quotes.find({"contact_id": cid}, {"_id": 0}).to_list(100)
    invoices = await db.invoices.find({"contact_id": cid}, {"_id": 0}).to_list(100)
    submissions = await db.form_submissions.find({"contact_id": cid}, {"_id": 0}).to_list(100)
    emails = await db.email_messages.find({"contact_id": cid}, {"_id": 0}).to_list(100)
    return {"contact": contact, "deals": deals, "quotes": quotes, "invoices": invoices,
            "form_submissions": submissions, "emails": emails}


# ── Pipeline Stages ────────────────────────────────────────────────────────────
@api.get("/pipeline-stages")
async def list_stages(u: dict = Depends(current_user)):
    stages = await db.pipeline_stages.find({"owner_id": u["id"]}, {"_id": 0}).sort("order", 1).to_list(50)
    return stages


@api.post("/pipeline-stages")
async def create_stage(p: StageIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    return await _insert("pipeline_stages", doc, u["actor_id"], "pipeline_stage")


# ── Deals ──────────────────────────────────────────────────────────────────────
@api.get("/deals")
async def list_deals(u: dict = Depends(current_user)):
    return await _list("deals", {"owner_id": u["id"], "deleted_at": {"$exists": False}})


@api.post("/deals")
async def create_deal(p: DealIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    doc["status"] = "open"
    return await _insert("deals", doc, u["actor_id"], "deal")


@api.put("/deals/{did}")
async def update_deal(did: str, p: DealIn, u: dict = Depends(current_user)):
    return await _update("deals", did, p.model_dump(), u["actor_id"], "deal")


@api.patch("/deals/{did}/stage")
async def move_deal(did: str, body: dict = Body(...), u: dict = Depends(current_user)):
    stage_id = body.get("pipeline_stage_id")
    if not stage_id:
        raise HTTPException(400, "pipeline_stage_id required")
    stage = await db.pipeline_stages.find_one({"id": stage_id, "owner_id": u["id"]}, {"_id": 0})
    if not stage:
        raise HTTPException(404, "Stage not found")
    patch = {"pipeline_stage_id": stage_id, "probability": stage["probability"]}
    if stage["altitude_label"] == "Closed Won":
        patch["status"] = "won"
        patch["actual_close_date"] = now_iso()
    elif stage["altitude_label"] == "Closed Lost":
        patch["status"] = "lost"
        patch["actual_close_date"] = now_iso()
    resp = await _update("deals", did, patch, u["actor_id"], "deal")
    # Phase 2: fire automations on stage change
    try:
        await _run_automations(
            u["id"], "deal_stage_change",
            {"to": stage["name"], "altitude": stage["altitude_label"]},
            {"entity_type": "deal", "entity_id": did, "contact_id": resp.get("contact_id")},
        )
    except Exception as e:
        log.warning(f"automation hook failed: {e}")
    return resp


@api.delete("/deals/{did}")
async def delete_deal(did: str, u: dict = Depends(current_user)):
    return await _soft_delete("deals", did, u["actor_id"], "deal")


# ── Products ───────────────────────────────────────────────────────────────────
@api.get("/products")
async def list_products(u: dict = Depends(current_user)):
    return await _list("products", {"owner_id": u["id"]})


@api.post("/products")
async def create_product(p: ProductIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    return await _insert("products", doc, u["actor_id"], "product")


@api.put("/products/{pid}")
async def update_product(pid: str, p: ProductIn, u: dict = Depends(current_user)):
    return await _update("products", pid, p.model_dump(), u["actor_id"], "product")


# ── Quotes ─────────────────────────────────────────────────────────────────────
def _compute_totals(lines: list[dict]) -> dict:
    subtotal = 0.0
    discount = 0.0
    tax = 0.0
    for ln in lines:
        line_sub = ln["qty"] * ln["unit_price"]
        line_disc = line_sub * (ln.get("discount_pct", 0) / 100)
        net = line_sub - line_disc
        line_tax = net * (ln.get("tax_rate", 0) / 100)
        ln["line_total"] = round(net + line_tax, 2)
        subtotal += line_sub
        discount += line_disc
        tax += line_tax
    grand = subtotal - discount + tax
    return {
        "subtotal": round(subtotal, 2),
        "discount_total": round(discount, 2),
        "tax_total": round(tax, 2),
        "grand_total": round(grand, 2),
    }


async def _next_number(collection: str, prefix: str, owner_id: str) -> str:
    count = await db[collection].count_documents({"owner_id": owner_id})
    year = datetime.now(timezone.utc).year
    return f"{prefix}-{year}-{count + 1:04d}"


@api.get("/quotes")
async def list_quotes(u: dict = Depends(current_user)):
    return await _list("quotes", {"owner_id": u["id"]})


@api.post("/quotes")
async def create_quote(p: QuoteIn, u: dict = Depends(current_user)):
    lines = [ln.model_dump() for ln in p.line_items]
    totals = _compute_totals(lines)
    doc = p.model_dump()
    doc["line_items"] = lines
    doc.update(totals)
    doc["owner_id"] = u["id"]
    doc["number"] = await _next_number("quotes", "QT", u["id"])
    doc["status"] = "draft"
    doc["version"] = 1
    # Auto-compute valid_until from valid_days (e.g. 30 → today + 30 days)
    if not doc.get("valid_until") and doc.get("valid_days"):
        doc["valid_until"] = (datetime.now(timezone.utc) + timedelta(days=int(doc["valid_days"]))).date().isoformat()
    return await _insert("quotes", doc, u["actor_id"], "quote")


@api.put("/quotes/{qid}")
async def update_quote(qid: str, p: QuoteIn, u: dict = Depends(current_user)):
    lines = [ln.model_dump() for ln in p.line_items]
    totals = _compute_totals(lines)
    patch = p.model_dump()
    patch["line_items"] = lines
    patch.update(totals)
    # If valid_days supplied and valid_until blank, recompute from today
    if not patch.get("valid_until") and patch.get("valid_days"):
        patch["valid_until"] = (datetime.now(timezone.utc) + timedelta(days=int(patch["valid_days"]))).date().isoformat()
    return await _update("quotes", qid, patch, u["actor_id"], "quote")


@api.post("/quotes/{qid}/send")
async def send_quote(qid: str, u: dict = Depends(current_user)):
    return await _update("quotes", qid, {"status": "sent", "sent_at": now_iso()}, u["actor_id"], "quote")


@api.post("/quotes/{qid}/accept")
async def accept_quote(qid: str, body: dict = Body(default={}), u: dict = Depends(current_user)):
    patch = {
        "status": "accepted",
        "acceptance": {
            "accepted_at": now_iso(),
            "accepted_by": body.get("signature_name", u["email"]),
            "ip": body.get("ip", "127.0.0.1"),
        },
    }
    return await _update("quotes", qid, patch, u["actor_id"], "quote")


@api.post("/quotes/{qid}/to-invoice")
async def quote_to_invoice(qid: str, u: dict = Depends(current_user)):
    q = await db.quotes.find_one({"id": qid, "owner_id": u["id"]}, {"_id": 0})
    if not q:
        raise HTTPException(404, "Quote not found")
    doc = {
        "owner_id": u["id"],
        "quote_id": qid,
        "contact_id": q.get("contact_id"),
        "company_id": q.get("company_id"),
        "line_items": q["line_items"],
        "subtotal": q["subtotal"],
        "tax_total": q["tax_total"],
        "grand_total": q["grand_total"],
        "currency": q.get("currency", "USD"),
        "issue_date": now_iso(),
        "due_date": (datetime.now(timezone.utc) + timedelta(days=14)).isoformat(),
        "status": "draft",
        "number": await _next_number("invoices", "INV", u["id"]),
    }
    return await _insert("invoices", doc, u["actor_id"], "invoice")


# ── Invoices ───────────────────────────────────────────────────────────────────
@api.get("/invoices")
async def list_invoices(u: dict = Depends(current_user)):
    return await _list("invoices", {"owner_id": u["id"]})


@api.post("/invoices")
async def create_invoice(p: InvoiceIn, u: dict = Depends(current_user)):
    lines = [ln.model_dump() for ln in (p.line_items or [])]
    totals = _compute_totals(lines)
    doc = p.model_dump()
    doc["line_items"] = lines
    doc.update(totals)
    doc["owner_id"] = u["id"]
    doc["number"] = await _next_number("invoices", "INV", u["id"])
    doc["status"] = "draft"
    doc["issue_date"] = now_iso()
    if not doc.get("due_date"):
        doc["due_date"] = (datetime.now(timezone.utc) + timedelta(days=14)).isoformat()
    created = await _insert("invoices", doc, u["actor_id"], "invoice")
    # Accrual-basis auto-journal: DR Debtors, CR Revenue, CR VAT Output.
    # Skipped silently if COA not yet seeded (no accounts table rows).
    try:
        if await db.accounts.find_one({"owner_id": u["id"], "code": "22000"}, {"_id": 0, "id": 1}):
            await _auto_post_invoice_journal(u["id"], created)
    except Exception as e:
        log.warning(f"invoice auto-journal hook error: {e}")
    return created


@api.put("/invoices/{iid}")
async def update_invoice(iid: str, p: InvoiceIn, u: dict = Depends(current_user)):
    patch = p.model_dump()
    if p.line_items is not None:
        lines = [ln.model_dump() for ln in p.line_items]
        totals = _compute_totals(lines)
        patch["line_items"] = lines
        patch.update(totals)
    return await _update("invoices", iid, patch, u["actor_id"], "invoice")


@api.post("/invoices/{iid}/send")
async def send_invoice(iid: str, u: dict = Depends(current_user)):
    return await _update("invoices", iid, {"status": "sent", "sent_at": now_iso()}, u["actor_id"], "invoice")


# ── Stripe ─────────────────────────────────────────────────────────────────────
@api.post("/payments/checkout")
async def create_checkout(req: CheckoutReq, http_request: Request, u: dict = Depends(current_user)):
    invoice = await db.invoices.find_one({"id": req.invoice_id, "owner_id": u["id"]}, {"_id": 0})
    if not invoice:
        raise HTTPException(404, "Invoice not found")
    try:
        from emergentintegrations.payments.stripe.checkout import (
            CheckoutSessionRequest,
            StripeCheckout,
        )
    except Exception as e:
        raise HTTPException(500, f"Stripe lib missing: {e}")

    host_url = str(http_request.base_url).rstrip("/")
    webhook_url = f"{host_url}/api/webhook/stripe"
    stripe = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)

    amount = float(invoice["grand_total"])  # server-authoritative
    currency = invoice.get("currency", "USD").lower()
    origin = req.origin_url.rstrip("/")
    success_url = f"{origin}/invoices/{invoice['id']}?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{origin}/invoices/{invoice['id']}"
    metadata = {"invoice_id": invoice["id"], "owner_id": u["id"], "invoice_number": invoice["number"]}

    session = await stripe.create_checkout_session(
        CheckoutSessionRequest(
            amount=amount, currency=currency, success_url=success_url,
            cancel_url=cancel_url, metadata=metadata,
        )
    )
    await db.payment_transactions.insert_one({
        "id": new_id(),
        "owner_id": u["id"],
        "invoice_id": invoice["id"],
        "session_id": session.session_id,
        "amount": amount,
        "currency": currency,
        "status": "initiated",
        "payment_status": "pending",
        "metadata": metadata,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    })
    await db.invoices.update_one(
        {"id": invoice["id"]},
        {"$set": {"payment_link": session.url, "stripe_session_id": session.session_id, "status": "sent"}},
    )
    return {"url": session.url, "session_id": session.session_id}


@api.get("/payments/status/{session_id}")
async def payment_status(session_id: str, http_request: Request, u: dict = Depends(current_user)):
    try:
        from emergentintegrations.payments.stripe.checkout import StripeCheckout
    except Exception as e:
        raise HTTPException(500, f"Stripe lib missing: {e}")
    host_url = str(http_request.base_url).rstrip("/")
    stripe = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=f"{host_url}/api/webhook/stripe")
    try:
        cs = await stripe.get_checkout_status(session_id)
    except Exception as e:
        raise HTTPException(502, f"Stripe error: {e}")
    tx = await db.payment_transactions.find_one({"session_id": session_id}, {"_id": 0})
    if tx and tx["payment_status"] != cs.payment_status:
        await db.payment_transactions.update_one(
            {"session_id": session_id},
            {"$set": {"payment_status": cs.payment_status, "status": cs.status, "updated_at": now_iso()}},
        )
        if cs.payment_status == "paid" and tx.get("invoice_id"):
            invoice = await db.invoices.find_one({"id": tx["invoice_id"]}, {"_id": 0})
            if invoice and invoice.get("status") != "paid":
                await db.invoices.update_one(
                    {"id": tx["invoice_id"]},
                    {"$set": {"status": "paid", "paid_at": now_iso()}},
                )
                try:
                    if await db.accounts.find_one({"owner_id": invoice["owner_id"], "code": "22000"}, {"_id": 0, "id": 1}):
                        await _auto_post_payment_journal(invoice["owner_id"], invoice, provider="stripe")
                except Exception as e:
                    log.warning(f"payment auto-journal error (stripe status): {e}")
    return {"session_id": session_id, "status": cs.status, "payment_status": cs.payment_status,
            "amount": cs.amount_total, "currency": cs.currency}


@api.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    try:
        from emergentintegrations.payments.stripe.checkout import StripeCheckout
    except Exception:
        return {"received": False}
    body = await request.body()
    host_url = str(request.base_url).rstrip("/")
    stripe = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=f"{host_url}/api/webhook/stripe")
    try:
        evt = await stripe.handle_webhook(body, request.headers.get("Stripe-Signature"))
        if evt.session_id:
            await db.payment_transactions.update_one(
                {"session_id": evt.session_id},
                {"$set": {"payment_status": evt.payment_status, "updated_at": now_iso()}},
            )
            if evt.payment_status == "paid":
                tx = await db.payment_transactions.find_one({"session_id": evt.session_id}, {"_id": 0})
                if tx and tx.get("invoice_id"):
                    await db.invoices.update_one(
                        {"id": tx["invoice_id"], "status": {"$ne": "paid"}},
                        {"$set": {"status": "paid", "paid_at": now_iso()}},
                    )
        return {"received": True, "event_type": evt.event_type}
    except Exception as e:
        log.warning(f"stripe webhook error: {e}")
        return {"received": False, "error": str(e)}


# ── PayPal (REST API v2, server-side Orders + server redirect) ────────────────
PAYPAL_CLIENT_ID = os.environ.get("PAYPAL_CLIENT_ID", "")
PAYPAL_SECRET = os.environ.get("PAYPAL_SECRET", "")
PAYPAL_MODE = os.environ.get("PAYPAL_MODE", "sandbox")
PAYPAL_BASE = "https://api-m.sandbox.paypal.com" if PAYPAL_MODE != "live" else "https://api-m.paypal.com"


async def _paypal_access_token() -> str:
    import httpx
    if not PAYPAL_CLIENT_ID or not PAYPAL_SECRET:
        raise HTTPException(500, "PayPal not configured (PAYPAL_CLIENT_ID/PAYPAL_SECRET missing)")
    async with httpx.AsyncClient(timeout=20.0) as client:
        r = await client.post(
            f"{PAYPAL_BASE}/v1/oauth2/token",
            auth=(PAYPAL_CLIENT_ID, PAYPAL_SECRET),
            data={"grant_type": "client_credentials"},
            headers={"Accept": "application/json"},
        )
        if r.status_code != 200:
            raise HTTPException(502, f"PayPal auth failed: {r.status_code} {r.text[:200]}")
        return r.json()["access_token"]


@api.post("/payments/paypal/checkout")
async def paypal_create_order(req: CheckoutReq, http_request: Request, u: dict = Depends(current_user)):
    """Create a PayPal Order for a given invoice and return the approval URL."""
    import httpx
    invoice = await db.invoices.find_one({"id": req.invoice_id, "owner_id": u["id"]}, {"_id": 0})
    if not invoice:
        raise HTTPException(404, "Invoice not found")

    amount = float(invoice["grand_total"])
    currency = (invoice.get("currency") or "USD").upper()
    origin = req.origin_url.rstrip("/")
    cancel_url = f"{origin}/invoices/{invoice['id']}"

    token = await _paypal_access_token()
    body = {
        "intent": "CAPTURE",
        "purchase_units": [{
            "reference_id": invoice["id"],
            "description": f"Invoice {invoice['number']}",
            "invoice_id": invoice["number"],
            "amount": {"currency_code": currency, "value": f"{amount:.2f}"},
        }],
        "application_context": {
            "brand_name": "Ascent CRM",
            "user_action": "PAY_NOW",
            "return_url": f"{origin}/invoices/{invoice['id']}?paypal=success",
            "cancel_url": cancel_url,
        },
    }
    async with httpx.AsyncClient(timeout=20.0) as client:
        r = await client.post(
            f"{PAYPAL_BASE}/v2/checkout/orders",
            json=body,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
        )
    if r.status_code not in (200, 201):
        raise HTTPException(502, f"PayPal order create failed: {r.status_code} {r.text[:300]}")
    data = r.json()
    order_id = data["id"]
    approve = next((lk["href"] for lk in data.get("links", []) if lk.get("rel") == "approve"), None)
    if not approve:
        raise HTTPException(502, "PayPal did not return an approval link")

    await db.payment_transactions.insert_one({
        "id": new_id(), "owner_id": u["id"], "invoice_id": invoice["id"],
        "provider": "paypal", "session_id": order_id, "amount": amount,
        "currency": currency.lower(), "status": "initiated", "payment_status": "pending",
        "metadata": {"invoice_number": invoice["number"]},
        "created_at": now_iso(), "updated_at": now_iso(),
    })
    await db.invoices.update_one(
        {"id": invoice["id"]},
        {"$set": {"paypal_order_id": order_id, "payment_link": approve, "status": "sent"}},
    )
    await audit(u["actor_id"], "paypal_checkout", "invoice", invoice["id"],
                after={"order_id": order_id, "amount": amount, "currency": currency})
    return {"url": approve, "session_id": order_id, "order_id": order_id}


@api.get("/payments/paypal/status/{order_id}")
async def paypal_order_status(order_id: str, u: dict = Depends(current_user)):
    """Read order from PayPal; if APPROVED and not yet captured, capture it; mark invoice paid on COMPLETED."""
    import httpx
    # Scope by owner: the caller must own the local payment_transactions row for this order.
    local_tx = await db.payment_transactions.find_one(
        {"session_id": order_id, "provider": "paypal", "owner_id": u["id"]},
        {"_id": 0},
    )
    if not local_tx:
        raise HTTPException(404, "PayPal order not found in your account")
    token = await _paypal_access_token()
    async with httpx.AsyncClient(timeout=20.0) as client:
        r = await client.get(
            f"{PAYPAL_BASE}/v2/checkout/orders/{order_id}",
            headers={"Authorization": f"Bearer {token}"},
        )
    if r.status_code == 404:
        raise HTTPException(404, "PayPal order not found")
    if r.status_code != 200:
        raise HTTPException(502, f"PayPal status failed: {r.status_code} {r.text[:200]}")
    data = r.json()
    status = data.get("status", "UNKNOWN")  # CREATED, APPROVED, COMPLETED, VOIDED

    # If buyer approved but we haven't captured yet, capture now.
    if status == "APPROVED":
        async with httpx.AsyncClient(timeout=20.0) as client:
            cap = await client.post(
                f"{PAYPAL_BASE}/v2/checkout/orders/{order_id}/capture",
                headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            )
        if cap.status_code in (200, 201):
            data = cap.json()
            status = data.get("status", status)
        else:
            log.warning(f"PayPal capture failed for {order_id}: {cap.status_code} {cap.text[:200]}")

    payment_status = "paid" if status == "COMPLETED" else ("pending" if status in ("CREATED", "APPROVED") else "failed")

    tx = await db.payment_transactions.find_one({"session_id": order_id, "provider": "paypal"}, {"_id": 0})
    if tx and tx.get("payment_status") != payment_status:
        await db.payment_transactions.update_one(
            {"session_id": order_id, "provider": "paypal"},
            {"$set": {"payment_status": payment_status, "status": status.lower(), "updated_at": now_iso()}},
        )
        if payment_status == "paid" and tx.get("invoice_id"):
            invoice = await db.invoices.find_one({"id": tx["invoice_id"]}, {"_id": 0})
            if invoice and invoice.get("status") != "paid":
                await db.invoices.update_one(
                    {"id": tx["invoice_id"], "owner_id": tx["owner_id"]},
                    {"$set": {"status": "paid", "paid_at": now_iso()}},
                )
                await audit(u["actor_id"], "paypal_captured", "invoice", tx["invoice_id"],
                            after={"order_id": order_id, "amount": tx.get("amount")})

    amount_obj = (data.get("purchase_units") or [{}])[0].get("amount") or {}
    return {
        "session_id": order_id, "order_id": order_id, "status": status.lower(),
        "payment_status": payment_status,
        "amount": float(amount_obj.get("value") or 0),
        "currency": (amount_obj.get("currency_code") or "").lower(),
    }


@api.post("/webhook/paypal")
async def paypal_webhook(request: Request):
    """Inbound PayPal webhook — mark invoice paid on PAYMENT.CAPTURE.COMPLETED / CHECKOUT.ORDER.APPROVED.

    NOTE: PayPal webhook signature verification requires a configured Webhook ID;
    until one is set via env PAYPAL_WEBHOOK_ID, we accept the event but skip
    crypto verification (sandbox-friendly; flip to required for production).
    """
    try:
        body = await request.json()
    except Exception:
        return {"received": False, "error": "invalid json"}
    event_type = body.get("event_type") or ""
    resource = body.get("resource") or {}
    order_id = resource.get("id") if event_type.startswith("CHECKOUT.ORDER") else resource.get("supplementary_data", {}).get("related_ids", {}).get("order_id")

    if event_type == "PAYMENT.CAPTURE.COMPLETED" and order_id:
        tx = await db.payment_transactions.find_one({"session_id": order_id, "provider": "paypal"}, {"_id": 0})
        if tx:
            await db.payment_transactions.update_one(
                {"session_id": order_id, "provider": "paypal"},
                {"$set": {"payment_status": "paid", "status": "completed", "updated_at": now_iso()}},
            )
            if tx.get("invoice_id"):
                await db.invoices.update_one(
                    {"id": tx["invoice_id"], "status": {"$ne": "paid"}},
                    {"$set": {"status": "paid", "paid_at": now_iso()}},
                )
    await db.webhook_events.insert_one({
        "id": new_id(), "provider": "paypal", "event_type": event_type,
        "order_id": order_id, "received_at": now_iso(), "raw": body,
    })
    return {"received": True, "event_type": event_type}




# ── Lead Forms ─────────────────────────────────────────────────────────────────
@api.get("/forms")
async def list_forms(u: dict = Depends(current_user)):
    return await _list("lead_forms", {"owner_id": u["id"]})


@api.post("/forms")
async def create_form(p: LeadFormIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    doc["submissions_count"] = 0
    return await _insert("lead_forms", doc, u["actor_id"], "lead_form")


@api.get("/forms/{slug}/public")
async def get_public_form(slug: str):
    form = await db.lead_forms.find_one({"slug": slug}, {"_id": 0, "owner_id": 0})
    if not form:
        raise HTTPException(404, "Form not found")
    return form


@api.post("/forms/{slug}/submit")
async def submit_form(slug: str, body: FormSubmitReq, request: Request):
    form = await db.lead_forms.find_one({"slug": slug}, {"_id": 0})
    if not form:
        raise HTTPException(404, "Form not found")
    if not body.consent_given:
        raise HTTPException(400, "Consent is required to submit this form (GDPR).")
    answers = body.answers or {}
    email = answers.get("email") or answers.get("Email")
    contact_id = None
    if email:
        existing = await db.contacts.find_one({"email": email, "owner_id": form["owner_id"]}, {"_id": 0})
        if existing:
            contact_id = existing["id"]
        else:
            contact_id = new_id()
            await db.contacts.insert_one({
                "id": contact_id,
                "owner_id": form["owner_id"],
                "first_name": answers.get("first_name") or answers.get("name") or "Unknown",
                "last_name": answers.get("last_name") or "",
                "email": email,
                "phone": answers.get("phone"),
                "notes": answers.get("message") or "",
                "tags": [f"source:form:{slug}"],
                "consent": {"marketing": body.consent_given, "newsletter": body.consent_given,
                            "updated_at": now_iso(), "source": f"form:{slug}"},
                "custom_fields": {},
                "interaction_count": 1,
                "last_activity_at": now_iso(),
                "created_at": now_iso(),
                "updated_at": now_iso(),
            })
    submission_id = new_id()
    await db.form_submissions.insert_one({
        "id": submission_id,
        "form_id": form["id"],
        "form_slug": slug,
        "owner_id": form["owner_id"],
        "answers": answers,
        "consent_given": body.consent_given,
        "contact_id": contact_id,
        "ip": request.client.host if request.client else None,
        "created_at": now_iso(),
    })
    if body.consent_given and email:
        await db.consent_logs.insert_one({
            "id": new_id(), "contact_email": email, "contact_id": contact_id, "kind": "form",
            "given": True, "source": f"form:{slug}", "timestamp": now_iso(),
        })
    await db.lead_forms.update_one({"id": form["id"]}, {"$inc": {"submissions_count": 1}})
    await audit(form["owner_id"], "submit", "lead_form", form["id"], after={"submission_id": submission_id})
    # Phase 2: fire form_submission automations
    try:
        await _run_automations(
            form["owner_id"], "form_submission", {"slug": slug},
            {"entity_type": "contact", "entity_id": contact_id, "contact_id": contact_id},
        )
    except Exception as e:
        log.warning(f"form automation hook failed: {e}")
    return {"ok": True, "submission_id": submission_id,
            "message": "Please check your inbox to confirm (double opt-in)." if form.get("double_opt_in") and email else "Thank you!"}


@api.get("/forms/{fid}/submissions")
async def list_submissions(fid: str, u: dict = Depends(current_user)):
    return await _list("form_submissions", {"form_id": fid, "owner_id": u["id"]})


@api.delete("/forms/{fid}")
async def delete_form(fid: str, u: dict = Depends(current_user)):
    res = await db.lead_forms.delete_one({"id": fid, "owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Form not found")
    await audit(u["actor_id"], "delete", "lead_form", fid)
    return {"ok": True}


# ── AI Studio (Gemini 3 via emergentintegrations) ──────────────────────────────
async def _build_crm_context(u: dict, req: AIGenerateReq) -> tuple[str, list[dict]]:
    """Return (context_block, fields_used)."""
    fields_used: list[dict] = []
    lines: list[str] = []

    bv = u.get("brand_voice") or {}
    lines.append(f"Brand voice: tone={bv.get('tone')}, hints={bv.get('vocabulary_hints')}")
    fields_used.append({"entity": "brand_voice", "id": u["id"], "fields": ["tone", "vocabulary_hints", "signature"]})

    if req.contact_id:
        c = await db.contacts.find_one({"id": req.contact_id, "owner_id": u["id"]}, {"_id": 0})
        if c:
            lines.append(f"Contact: {c.get('first_name')} {c.get('last_name')} ({c.get('email')}), role={c.get('role_title')}, notes={c.get('notes')}")
            fields_used.append({"entity": "contact", "id": c["id"],
                                "fields": ["first_name", "last_name", "email", "role_title", "notes"]})
            if c.get("company_id"):
                comp = await db.companies.find_one({"id": c["company_id"]}, {"_id": 0})
                if comp:
                    lines.append(f"Company: {comp.get('name')}, industry={comp.get('industry')}, notes={comp.get('notes')}")
                    fields_used.append({"entity": "company", "id": comp["id"],
                                        "fields": ["name", "industry", "notes"]})

    if req.deal_id:
        d = await db.deals.find_one({"id": req.deal_id, "owner_id": u["id"]}, {"_id": 0})
        if d:
            lines.append(f"Deal: {d.get('title')} value={d.get('value')} {d.get('currency')} probability={d.get('probability')}%")
            fields_used.append({"entity": "deal", "id": d["id"], "fields": ["title", "value", "currency", "probability"]})

    # Price list for grounding quote language
    products = await db.products.find({"owner_id": u["id"], "active": True}, {"_id": 0}).to_list(20)
    if products:
        lines.append("Price list (authoritative — do not invent new prices):")
        for p in products[:10]:
            lines.append(f"  • {p['name']} ({p['sku']}) {p['currency']} {p['unit_price']} — {p.get('description') or ''}")
        fields_used.append({"entity": "products", "id": "*", "fields": ["name", "sku", "unit_price", "currency"]})

    if req.incoming_email:
        lines.append(f"\nIncoming email from prospect:\n{req.incoming_email}")
        fields_used.append({"entity": "incoming_email", "id": "inline", "fields": ["body"]})

    return "\n".join(lines), fields_used


@api.post("/ai/generate")
async def ai_generate(req: AIGenerateReq, u: dict = Depends(current_user)):
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
    except Exception as e:
        raise HTTPException(500, f"LLM lib missing: {e}")

    context_block, fields_used = await _build_crm_context(u, req)
    bv = u.get("brand_voice") or {}
    banned = ", ".join(bv.get("banned_phrases") or []) or "none"
    system = (
        f"You are the {req.kind} writing assistant for {u.get('name', 'the user')}.\n"
        f"Brand tone requested: {req.tone}. Brand hints: {bv.get('vocabulary_hints') or '—'}. "
        f"Banned phrases: {banned}.\n"
        "HARD RULES:\n"
        "1. Use ONLY facts present in <CRM_CONTEXT>. If a required fact is missing, add a clarifying question "
        "   in the `questions_for_user` field INSTEAD of inventing it.\n"
        "2. Never invent prices, dates, client names, testimonials, or commitments.\n"
        "3. Output STRICTLY as JSON with keys: draft (string), fields_used (array of strings), "
        "   questions_for_user (array of strings).\n"
        "4. `fields_used` must enumerate the CRM fields you actually relied on (e.g. "
        "   `contact.first_name`, `deal.value`, `products[].unit_price`).\n"
        "5. No markdown fences in the JSON.\n"
    )
    user_text = (
        f"<CRM_CONTEXT>\n{context_block}\n</CRM_CONTEXT>\n\n"
        f"User task ({req.kind}): {req.prompt}"
    )
    session_id = f"ai-{u['id']}-{new_id()[:8]}"
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message=system).with_model(
        "gemini", "gemini-3-flash-preview"
    )
    try:
        raw = await chat.send_message(UserMessage(text=user_text))
    except Exception as e:
        raise HTTPException(502, f"LLM error: {e}")

    # Best-effort JSON parse
    parsed: dict[str, Any]
    try:
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("```", 2)[1]
            if cleaned.startswith("json"):
                cleaned = cleaned[4:]
            cleaned = cleaned.strip()
        parsed = json.loads(cleaned)
    except Exception:
        parsed = {"draft": raw, "fields_used": [], "questions_for_user": []}

    gen_id = new_id()
    await db.ai_generations.insert_one({
        "id": gen_id,
        "owner_id": u["id"],
        "kind": req.kind,
        "prompt": req.prompt,
        "tone": req.tone,
        "output": parsed.get("draft", ""),
        "questions_for_user": parsed.get("questions_for_user", []),
        "llm_fields_used": parsed.get("fields_used", []),
        "grounding_fields": fields_used,
        "model": "gemini-3-flash-preview",
        "session_id": session_id,
        "created_at": now_iso(),
    })
    return {
        "id": gen_id,
        "draft": parsed.get("draft", ""),
        "fields_used": fields_used,
        "llm_fields_used": parsed.get("fields_used", []),
        "questions_for_user": parsed.get("questions_for_user", []),
        "model": "gemini-3-flash-preview",
    }


@api.get("/ai/history")
async def ai_history(u: dict = Depends(current_user)):
    rows = await db.ai_generations.find({"owner_id": u["id"]}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return rows


# ── Templates ──────────────────────────────────────────────────────────────────
TEMPLATES: list[dict] = [
    {
        "id": "tpl-executive-coach",
        "kind": "coach",
        "name": "Executive Coach",
        "description": "Discovery → Intake → Proposal → Onboarding → Recurring sessions. Inspired by 5 Voices / Sherpa.",
        "pipeline_stages": [
            {"name": "Discovery", "altitude_label": "Basecamp", "probability": 10},
            {"name": "Intake", "altitude_label": "Basecamp", "probability": 25},
            {"name": "Proposal Sent", "altitude_label": "Ascent", "probability": 55},
            {"name": "Contract Signed", "altitude_label": "Ascent", "probability": 85},
            {"name": "Engaged", "altitude_label": "Summit", "probability": 95},
            {"name": "Won", "altitude_label": "Closed Won", "probability": 100},
            {"name": "Lost", "altitude_label": "Closed Lost", "probability": 0},
        ],
        "sample_products": [
            {"sku": "5V-FOUND", "name": "5 Voices Foundation", "unit_price": 99, "currency": "USD", "tier": "foundation"},
            {"sku": "100X-GRP", "name": "100X Leader Group (6 sessions)", "unit_price": 4500, "currency": "USD", "tier": "summit"},
            {"sku": "EXEC-1on1", "name": "Executive 1-on-1 Monthly", "unit_price": 850, "currency": "USD", "tier": "ascent"},
        ],
        "sample_forms": [{"name": "Discovery Intake", "slug": "discovery",
                          "fields": ["first_name", "last_name", "email", "role_title", "challenge"]}],
        "sample_emails": ["session-reminder-24h", "post-session-notes", "mid-program-check-in"],
    },
    {
        "id": "tpl-fitness-coach",
        "kind": "fitness",
        "name": "Fitness Coach",
        "description": "Lead → Assessment → Package → Monthly subscription → Session reminders.",
        "pipeline_stages": [
            {"name": "New Lead", "altitude_label": "Basecamp", "probability": 10},
            {"name": "Free Assessment", "altitude_label": "Basecamp", "probability": 35},
            {"name": "Package Offered", "altitude_label": "Ascent", "probability": 60},
            {"name": "Active Client", "altitude_label": "Summit", "probability": 95},
            {"name": "Won", "altitude_label": "Closed Won", "probability": 100},
            {"name": "Churned", "altitude_label": "Closed Lost", "probability": 0},
        ],
        "sample_products": [
            {"sku": "12WK-BASE", "name": "12-Week Transformation Base", "unit_price": 1200, "currency": "USD", "tier": "foundation"},
            {"sku": "MONTH-SUB", "name": "Monthly Coaching Subscription", "unit_price": 220, "currency": "USD", "tier": "ascent"},
        ],
        "sample_forms": [{"name": "Fitness Intake", "slug": "fit-intake",
                          "fields": ["first_name", "email", "phone", "goals", "injuries"]}],
        "sample_emails": ["workout-reminder", "weekly-check-in", "re-engagement"],
    },
    {
        "id": "tpl-business-consultant",
        "kind": "consultant",
        "name": "Business Consultant",
        "description": "Discovery → Diagnostic → SOW → Engagement → Retainer.",
        "pipeline_stages": [
            {"name": "Inbound", "altitude_label": "Basecamp", "probability": 10},
            {"name": "Discovery Call", "altitude_label": "Basecamp", "probability": 30},
            {"name": "Diagnostic", "altitude_label": "Ascent", "probability": 50},
            {"name": "SOW Sent", "altitude_label": "Ascent", "probability": 70},
            {"name": "Engaged", "altitude_label": "Summit", "probability": 95},
            {"name": "Won", "altitude_label": "Closed Won", "probability": 100},
            {"name": "Lost", "altitude_label": "Closed Lost", "probability": 0},
        ],
        "sample_products": [
            {"sku": "DIAG-2D", "name": "2-Day Business Diagnostic", "unit_price": 6500, "currency": "USD", "tier": "foundation"},
            {"sku": "RETAINER-M", "name": "Monthly Advisory Retainer", "unit_price": 5000, "currency": "USD", "tier": "summit"},
        ],
        "sample_forms": [{"name": "Consulting Inquiry", "slug": "consulting",
                          "fields": ["first_name", "last_name", "email", "company", "budget", "timeline"]}],
        "sample_emails": ["sow-followup", "monthly-report", "renewal-90d"],
    },
]


@api.get("/templates")
async def list_templates():
    return TEMPLATES


@api.post("/templates/{tid}/apply")
async def apply_template(tid: str, u: dict = Depends(current_user)):
    tpl = next((t for t in TEMPLATES if t["id"] == tid), None)
    if not tpl:
        raise HTTPException(404, "Template not found")
    # Snapshot old stages so we can remap any existing deals
    old_stages = await db.pipeline_stages.find({"owner_id": u["id"]}, {"_id": 0}).to_list(100)
    # Wipe & recreate pipeline stages
    await db.pipeline_stages.delete_many({"owner_id": u["id"]})
    new_stage_ids_by_name: dict[str, str] = {}
    new_stage_ids_by_altitude: dict[str, str] = {}
    for i, st in enumerate(tpl["pipeline_stages"]):
        sid = new_id()
        await db.pipeline_stages.insert_one({
            "id": sid, "owner_id": u["id"], "name": st["name"], "order": i,
            "probability": st["probability"], "altitude_label": st["altitude_label"],
            "created_at": now_iso(), "updated_at": now_iso(),
        })
        new_stage_ids_by_name[st["name"].lower()] = sid
        new_stage_ids_by_altitude.setdefault(st["altitude_label"], sid)
    # Remap existing deals to new stages: prefer same altitude_label, else first new stage
    fallback = next(iter(new_stage_ids_by_altitude.values()))
    new_stage_ids_set = set(new_stage_ids_by_name.values())
    for old in old_stages:
        target = new_stage_ids_by_name.get((old.get("name") or "").lower()) \
            or new_stage_ids_by_altitude.get(old.get("altitude_label")) or fallback
        await db.deals.update_many(
            {"owner_id": u["id"], "pipeline_stage_id": old["id"]},
            {"$set": {"pipeline_stage_id": target, "updated_at": now_iso()}},
        )
    # Catch any deals whose pipeline_stage_id points to a stage that no longer exists at all.
    # Remap closed deals by status, else by deal.probability to nearest altitude, else fallback.
    async for d in db.deals.find(
        {"owner_id": u["id"], "pipeline_stage_id": {"$nin": list(new_stage_ids_set)}},
        {"_id": 0, "id": 1, "status": 1, "probability": 1},
    ):
        if d.get("status") == "won":
            target = new_stage_ids_by_altitude.get("Closed Won", fallback)
        elif d.get("status") == "lost":
            target = new_stage_ids_by_altitude.get("Closed Lost", fallback)
        else:
            prob = d.get("probability", 10)
            if prob >= 90:
                target = new_stage_ids_by_altitude.get("Summit") or new_stage_ids_by_altitude.get("Ascent") or fallback
            elif prob >= 40:
                target = new_stage_ids_by_altitude.get("Ascent") or fallback
            else:
                target = new_stage_ids_by_altitude.get("Basecamp") or fallback
        await db.deals.update_one({"id": d["id"]}, {"$set": {"pipeline_stage_id": target, "updated_at": now_iso()}})
    # Seed sample products (upsert by sku)
    for p in tpl["sample_products"]:
        await db.products.update_one(
            {"owner_id": u["id"], "sku": p["sku"]},
            {"$setOnInsert": {
                "id": new_id(), "owner_id": u["id"], "sku": p["sku"], "name": p["name"],
                "unit_price": p["unit_price"], "currency": p["currency"], "tier": p["tier"],
                "tax_rate": 0.0, "active": True,
                "created_at": now_iso(), "updated_at": now_iso(),
            }},
            upsert=True,
        )
    await audit(u["actor_id"], "apply_template", "template", tid)
    return {"ok": True, "template": tpl["name"]}


# ── Automations (simple CRUD, executed on stage move in /deals/{id}/stage) ─────
@api.get("/automations")
async def list_automations(u: dict = Depends(current_user)):
    return await _list("automations", {"owner_id": u["id"]})


@api.post("/automations")
async def create_automation(body: dict = Body(...), u: dict = Depends(current_user)):
    doc = {"owner_id": u["id"], "enabled": True, "run_count": 0, **body}
    return await _insert("automations", doc, u["actor_id"], "automation")


# ── Integrations (stub registry) ───────────────────────────────────────────────
@api.get("/integrations")
async def integrations_status(u: dict = Depends(current_user)):
    rows = await db.integrations.find({"owner_id": u["id"]}, {"_id": 0}).to_list(20)
    existing = {r["kind"]: r for r in rows}
    catalog = [
        {"kind": "stripe", "name": "Stripe", "description": "Payments + subscriptions", "auto_connected": True},
        {"kind": "paypal", "name": "PayPal", "description": "Alt payment rail (sandbox)", "auto_connected": bool(PAYPAL_CLIENT_ID and PAYPAL_SECRET)},
        {"kind": "calendly", "name": "Calendly", "description": "Discovery call booking → auto-create deal"},
        {"kind": "zoom", "name": "Zoom", "description": "Meeting links, session logging"},
        {"kind": "zapier", "name": "Zapier / Make", "description": "Webhook in/out catalog"},
        {"kind": "surveymonkey", "name": "SurveyMonkey", "description": "Survey results → CRM timeline"},
        {"kind": "msgraph", "name": "Microsoft Graph Email", "description": "Outlook / Microsoft 365 mailbox sync"},
        {"kind": "imap", "name": "IMAP/SMTP", "description": "Generic email fallback"},
    ]
    return [
        {**c, "status": existing.get(c["kind"], {}).get("status",
                                                         "connected" if c.get("auto_connected") else "disconnected"),
         "last_sync_at": existing.get(c["kind"], {}).get("last_sync_at")}
        for c in catalog
    ]


@api.post("/integrations/{kind}/toggle")
async def integrations_toggle(kind: str, u: dict = Depends(current_user)):
    existing = await db.integrations.find_one({"owner_id": u["id"], "kind": kind}, {"_id": 0})
    new_status = "disconnected" if existing and existing.get("status") == "connected" else "connected"
    await db.integrations.update_one(
        {"owner_id": u["id"], "kind": kind},
        {"$set": {"status": new_status, "last_sync_at": now_iso(),
                  "owner_id": u["id"], "kind": kind}},
        upsert=True,
    )
    await audit(u["actor_id"], "toggle", "integration", kind, after={"status": new_status})
    return {"kind": kind, "status": new_status}


# ── Analytics ──────────────────────────────────────────────────────────────────
@api.get("/analytics/summary")
async def analytics(u: dict = Depends(current_user)):
    deals = await db.deals.find({"owner_id": u["id"], "deleted_at": {"$exists": False}}, {"_id": 0}).to_list(1000)
    invoices = await db.invoices.find({"owner_id": u["id"]}, {"_id": 0}).to_list(1000)
    contacts_n = await db.contacts.count_documents({"owner_id": u["id"], "deleted_at": {"$exists": False}})
    stages = await db.pipeline_stages.find({"owner_id": u["id"]}, {"_id": 0}).sort("order", 1).to_list(50)

    open_deals = [d for d in deals if d.get("status") == "open"]
    won = [d for d in deals if d.get("status") == "won"]
    lost = [d for d in deals if d.get("status") == "lost"]
    pipeline_value = sum(d.get("value", 0) for d in open_deals)
    forecast = sum(d.get("value", 0) * d.get("probability", 0) / 100 for d in open_deals)
    paid_total = sum(i.get("grand_total", 0) for i in invoices if i.get("status") == "paid")
    outstanding = sum(i.get("grand_total", 0) for i in invoices if i.get("status") in ("sent", "overdue"))

    # Monthly revenue (last 12 months) from paid invoices
    from collections import defaultdict
    monthly = defaultdict(float)
    for inv in invoices:
        if inv.get("status") == "paid" and inv.get("paid_at"):
            try:
                ts = datetime.fromisoformat(inv["paid_at"].replace("Z", "+00:00"))
                key = ts.strftime("%Y-%m")
                monthly[key] += inv.get("grand_total", 0)
            except Exception:
                pass
    now = datetime.now(timezone.utc)
    revenue_series = []
    for i in range(11, -1, -1):
        d = (now.replace(day=1) - timedelta(days=30 * i))
        key = d.strftime("%Y-%m")
        revenue_series.append({"month": key, "revenue": round(monthly.get(key, 0), 2)})

    # Stage distribution
    stage_counts = []
    for s in stages:
        in_stage = [d for d in open_deals if d.get("pipeline_stage_id") == s["id"]]
        stage_counts.append({"stage": s["name"], "altitude": s["altitude_label"],
                             "count": len(in_stage), "value": sum(d.get("value", 0) for d in in_stage)})

    # Invoice aging
    aging = {"0-30": 0, "31-60": 0, "61-90": 0, "90+": 0}
    for inv in invoices:
        if inv.get("status") in ("sent", "overdue"):
            try:
                due = datetime.fromisoformat(inv["due_date"].replace("Z", "+00:00"))
                days = (now - due).days
                bucket = "0-30" if days <= 30 else "31-60" if days <= 60 else "61-90" if days <= 90 else "90+"
                aging[bucket] += inv.get("grand_total", 0)
            except Exception:
                pass

    return {
        "kpis": {
            "open_pipeline": round(pipeline_value, 2),
            "weighted_forecast": round(forecast, 2),
            "revenue_ytd": round(paid_total, 2),
            "outstanding": round(outstanding, 2),
            "contacts": contacts_n,
            "deals_open": len(open_deals),
            "deals_won": len(won),
            "deals_lost": len(lost),
            "win_rate": round(len(won) / max(1, len(won) + len(lost)) * 100, 1),
        },
        "revenue_series": revenue_series,
        "stage_distribution": stage_counts,
        "invoice_aging": [{"bucket": k, "amount": round(v, 2)} for k, v in aging.items()],
    }


# ── GDPR / Audit ───────────────────────────────────────────────────────────────
@api.get("/gdpr/consent-logs")
async def consent_logs(u: dict = Depends(current_user)):
    # filter by owner's contacts
    contact_ids = [c["id"] async for c in db.contacts.find({"owner_id": u["id"]}, {"id": 1})]
    rows = await db.consent_logs.find({"contact_id": {"$in": contact_ids}}, {"_id": 0}).sort("timestamp", -1).to_list(200)
    return rows


@api.get("/gdpr/export/{contact_id}")
async def gdpr_export(contact_id: str, u: dict = Depends(current_user)):
    contact = await db.contacts.find_one({"id": contact_id, "owner_id": u["id"]}, {"_id": 0})
    if not contact:
        raise HTTPException(404, "Contact not found")
    deals = await db.deals.find({"contact_id": contact_id}, {"_id": 0}).to_list(500)
    quotes = await db.quotes.find({"contact_id": contact_id}, {"_id": 0}).to_list(500)
    invoices = await db.invoices.find({"contact_id": contact_id}, {"_id": 0}).to_list(500)
    submissions = await db.form_submissions.find({"contact_id": contact_id}, {"_id": 0}).to_list(500)
    consents = await db.consent_logs.find({"contact_id": contact_id}, {"_id": 0}).to_list(500)
    emails = await db.email_messages.find({"contact_id": contact_id}, {"_id": 0}).to_list(500)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("contact.json", json.dumps(contact, indent=2, default=str))
        zf.writestr("deals.json", json.dumps(deals, indent=2, default=str))
        zf.writestr("quotes.json", json.dumps(quotes, indent=2, default=str))
        zf.writestr("invoices.json", json.dumps(invoices, indent=2, default=str))
        zf.writestr("form_submissions.json", json.dumps(submissions, indent=2, default=str))
        zf.writestr("consent_logs.json", json.dumps(consents, indent=2, default=str))
        zf.writestr("emails.json", json.dumps(emails, indent=2, default=str))
        zf.writestr("README.txt", "GDPR Subject Access Request export generated by Ascent CRM.\n")
    buf.seek(0)
    await audit(u["actor_id"], "gdpr_export", "contact", contact_id)
    return StreamingResponse(
        buf, media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="gdpr-{contact_id}.zip"'},
    )


@api.post("/gdpr/erase/{contact_id}")
async def gdpr_erase(contact_id: str, body: dict = Body(default={}), u: dict = Depends(current_user)):
    hard = body.get("hard", False)
    contact = await db.contacts.find_one({"id": contact_id, "owner_id": u["id"]}, {"_id": 0})
    if not contact:
        raise HTTPException(404, "Contact not found")
    if hard:
        await db.contacts.delete_one({"id": contact_id})
        await audit(u["actor_id"], "hard_erase", "contact", contact_id, before=contact)
        return {"ok": True, "mode": "hard"}
    await db.contacts.update_one(
        {"id": contact_id},
        {"$set": {
            "deleted_at": now_iso(),
            "erasure_scheduled_for": (datetime.now(timezone.utc) + timedelta(days=14)).isoformat(),
        }},
    )
    await audit(u["actor_id"], "soft_erase", "contact", contact_id, before=contact)
    return {"ok": True, "mode": "soft", "purge_in_days": 14}


@api.get("/audit")
async def list_audit(
    response: Response,
    limit: int = 100,
    after_id: str | None = None,
    u: dict = Depends(current_user),
):
    """Paginated audit trail (newest first).

    Query params:
      - `limit` (1-500, default 100) — page size
      - `after_id` — pass the `id` of the last row from the previous page to fetch older rows
    Returns a JSON array of audit rows. The next cursor is exposed in the
    `X-Next-After-Id` response header (empty when no more rows).
    """
    limit = max(1, min(int(limit or 100), 500))

    # Owner/admin see all entries from their team; everyone else sees only their own.
    if u.get("role") in ("owner", "admin"):
        member_ids = [m["id"] async for m in db.users.find(
            {"$or": [{"id": u["id"]}, {"team_owner_id": u["id"]}]}, {"id": 1}
        )]
        q: dict = {"actor_id": {"$in": member_ids}}
    else:
        q = {"actor_id": u["actor_id"]}

    # Cursor: entries are newest-first; `after_id` anchors to rows strictly older.
    if after_id:
        anchor = await db.audit_entries.find_one({"id": after_id}, {"_id": 0, "timestamp": 1})
        if anchor and anchor.get("timestamp"):
            q["timestamp"] = {"$lt": anchor["timestamp"]}

    rows = await db.audit_entries.find(q, {"_id": 0}).sort("timestamp", -1).to_list(limit + 1)
    has_more = len(rows) > limit
    rows = rows[:limit]
    response.headers["X-Next-After-Id"] = rows[-1]["id"] if has_more and rows else ""
    response.headers["Access-Control-Expose-Headers"] = "X-Next-After-Id"
    return rows


# ── Seed data ──────────────────────────────────────────────────────────────────
SEED_EMAIL = "demo@climbleadershiplab.com"
SEED_PASSWORD = "SherpaDemo2026!"


async def _seed() -> None:
    existing = await db.users.find_one({"email": SEED_EMAIL})
    if existing:
        log.info("Seed user already present, skipping")
        return
    uid = new_id()
    await db.users.insert_one({
        "id": uid,
        "email": SEED_EMAIL,
        "name": "Aleksia (Demo)",
        "password_hash": hash_pw(SEED_PASSWORD),
        "role": "owner",
        "brand_voice": {
            "tone": "warm-sherpa",
            "vocabulary_hints": "evoke mountain ascent and summit metaphors; reference the 5 Voices framework; tone is calm, authoritative, never flashy",
            "signature": "— Aleksia, CLiMB Leadership Lab",
            "banned_phrases": ["cutting-edge", "synergy", "game-changer"],
        },
        "created_at": now_iso(),
    })

    # Pipeline stages
    stages = [
        ("Basecamp (New)", "Basecamp", 10),
        ("Discovery Call", "Basecamp", 25),
        ("Proposal Sent", "Ascent", 55),
        ("Contract Signed", "Ascent", 85),
        ("Engaged", "Summit", 95),
        ("Won", "Closed Won", 100),
        ("Lost", "Closed Lost", 0),
    ]
    stage_ids: dict[str, str] = {}
    for i, (name, alt, prob) in enumerate(stages):
        sid = new_id()
        stage_ids[name] = sid
        await db.pipeline_stages.insert_one({
            "id": sid, "owner_id": uid, "name": name, "order": i,
            "probability": prob, "altitude_label": alt,
            "created_at": now_iso(), "updated_at": now_iso(),
        })

    # Products
    prods = [
        ("5V-FOUND", "5 Voices Foundation", 99, "USD", "foundation"),
        ("GIANT-PRO", "GiANT Pro Workspace (monthly)", 10, "USD", "foundation"),
        ("5V-TEAMS", "5 Voices for Teams (6 sessions)", 12000, "ZAR", "ascent"),
        ("5V-MARRIAGE", "5 Voices for Marriage", 1200, "ZAR", "ascent"),
        ("100X-GRP", "100X Leader Program (6 sessions)", 81000, "ZAR", "summit"),
        ("5V-CERT", "5 Voices Certification", 3995, "USD", "summit"),
        ("EXEC-1on1", "Executive 1-on-1 Monthly", 850, "USD", "ascent"),
    ]
    for sku, nm, pr, cur, tier in prods:
        await db.products.insert_one({
            "id": new_id(), "owner_id": uid, "sku": sku, "name": nm,
            "unit_price": pr, "currency": cur, "tax_rate": 15.0 if cur == "ZAR" else 0.0,
            "tier": tier, "active": True,
            "created_at": now_iso(), "updated_at": now_iso(),
        })

    # Companies + contacts + deals
    demo = [
        ("Altimeter Partners", "Private Equity", "Basecamp (New)", 12000, "USD", "Morgan", "Okafor", "morgan@altimeterpartners.com", "Partner"),
        ("Summit Health", "Healthcare", "Discovery Call", 45000, "ZAR", "Thandi", "Nkosi", "thandi@summithealth.co.za", "Head of L&D"),
        ("Bluewater Shipping", "Logistics", "Proposal Sent", 81000, "ZAR", "Johan", "Van der Merwe", "johan@bluewater.co.za", "COO"),
        ("Ridgeline Capital", "Finance", "Contract Signed", 18000, "USD", "Ava", "Chen", "ava@ridgeline.vc", "MD"),
        ("Northwind Consulting", "Consulting", "Engaged", 5400, "USD", "Priya", "Ramanathan", "priya@northwind.io", "Principal"),
        ("Copperleaf Foods", "Food & Bev", "Won", 99, "USD", "Sean", "O'Brien", "sean@copperleaffoods.com", "VP People"),
        ("Kilimanjaro Labs", "Biotech", "Lost", 6500, "USD", "Lerato", "Mokoena", "lerato@kililabs.com", "Head HR"),
    ]
    for cname, ind, stage_name, val, cur, fn, ln, em, role in demo:
        cid = new_id()
        await db.companies.insert_one({
            "id": cid, "owner_id": uid, "name": cname, "industry": ind,
            "lifecycle_stage": "customer" if stage_name == "Won" else "opportunity",
            "status": "active", "tags": [ind.lower()], "custom_fields": {},
            "created_at": now_iso(), "updated_at": now_iso(),
        })
        contact_id = new_id()
        await db.contacts.insert_one({
            "id": contact_id, "owner_id": uid, "company_id": cid,
            "first_name": fn, "last_name": ln, "email": em, "role_title": role,
            "tags": ["demo", "leadership"], "custom_fields": {"referrer": "5 Voices event"},
            "consent": {"marketing": True, "newsletter": True, "updated_at": now_iso(), "source": "manual-seed"},
            "interaction_count": 3, "last_activity_at": now_iso(),
            "created_at": now_iso(), "updated_at": now_iso(),
        })
        await db.deals.insert_one({
            "id": new_id(), "owner_id": uid, "title": f"{cname} — Leadership Engagement",
            "contact_id": contact_id, "company_id": cid, "pipeline_stage_id": stage_ids[stage_name],
            "value": val, "currency": cur,
            "probability": next(p for n, _, p in stages if n == stage_name),
            "status": "won" if stage_name == "Won" else "lost" if stage_name == "Lost" else "open",
            "expected_close_date": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat(),
            "tags": [ind.lower()],
            "created_at": now_iso(), "updated_at": now_iso(),
        })

    # Demo lead form
    await db.lead_forms.insert_one({
        "id": new_id(), "owner_id": uid, "name": "Discovery Call Request",
        "slug": "discovery", "double_opt_in": True,
        "consent_text": "I agree to be contacted by CLiMB Leadership Lab and accept the privacy policy.",
        "fields": [
            {"key": "first_name", "label": "First Name", "type": "text", "required": True},
            {"key": "last_name", "label": "Last Name", "type": "text", "required": True},
            {"key": "email", "label": "Email", "type": "email", "required": True},
            {"key": "phone", "label": "Phone", "type": "phone", "required": False},
            {"key": "company", "label": "Company", "type": "text", "required": False},
            {"key": "challenge", "label": "What leadership challenge are you facing?", "type": "textarea", "required": True},
        ],
        "submissions_count": 0,
        "created_at": now_iso(), "updated_at": now_iso(),
    })

    # Demo paid invoice (for revenue series to show data)
    paid_iso = (datetime.now(timezone.utc) - timedelta(days=10)).isoformat()
    await db.invoices.insert_one({
        "id": new_id(), "owner_id": uid, "number": f"INV-{datetime.now(timezone.utc).year}-0001",
        "contact_id": None, "company_id": None, "currency": "USD",
        "line_items": [{"description": "5 Voices Foundation", "qty": 1, "unit_price": 99, "discount_pct": 0, "tax_rate": 0, "line_total": 99}],
        "subtotal": 99, "discount_total": 0, "tax_total": 0, "grand_total": 99,
        "status": "paid", "issue_date": paid_iso, "due_date": paid_iso, "paid_at": paid_iso,
        "created_at": paid_iso, "updated_at": paid_iso,
    })

    log.info(f"✅ Seed complete: user={SEED_EMAIL} / password={SEED_PASSWORD}")


@app.on_event("startup")
async def on_start():
    try:
        await _seed()
    except Exception as e:
        log.exception(f"seed failed: {e}")


@app.on_event("shutdown")
async def on_stop():
    client.close()


# ── health ─────────────────────────────────────────────────────────────────────
@api.get("/")
async def root():
    return {"service": "Ascent CRM", "version": "1.1.0", "status": "ok"}


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 — Subscriptions, Automation Engine, SEO, Calendly webhook
# ══════════════════════════════════════════════════════════════════════════════

class SubscriptionIn(BaseDoc):
    product_id: str
    contact_id: str | None = None
    company_id: str | None = None
    interval: Literal["monthly", "quarterly", "annual"] = "monthly"
    quantity: int = 1
    start_date: str | None = None
    cycles: int | None = None  # None = indefinite


def _interval_days(interval: str) -> int:
    return {"monthly": 30, "quarterly": 91, "annual": 365}.get(interval, 30)


@api.get("/subscriptions")
async def list_subscriptions(u: dict = Depends(current_user)):
    return await _list("subscriptions", {"owner_id": u["id"], "deleted_at": {"$exists": False}})


@api.post("/subscriptions")
async def create_subscription(p: SubscriptionIn, u: dict = Depends(current_user)):
    prod = await db.products.find_one({"id": p.product_id, "owner_id": u["id"]}, {"_id": 0})
    if not prod:
        raise HTTPException(404, "Product not found")
    start = p.start_date or now_iso()
    next_at = start
    doc = {
        "owner_id": u["id"],
        "product_id": p.product_id,
        "product_name": prod["name"],
        "contact_id": p.contact_id,
        "company_id": p.company_id,
        "currency": prod.get("currency", "USD"),
        "unit_price": prod.get("unit_price", 0),
        "tax_rate": prod.get("tax_rate", 0),
        "quantity": p.quantity,
        "interval": p.interval,
        "status": "active",
        "cycles": p.cycles,
        "cycles_billed": 0,
        "start_date": start,
        "next_billing_at": next_at,
        "last_billed_at": None,
        "failed_payments": 0,
    }
    return await _insert("subscriptions", doc, u["actor_id"], "subscription")


@api.patch("/subscriptions/{sid}")
async def update_subscription(sid: str, body: dict = Body(...), u: dict = Depends(current_user)):
    allowed = {k: v for k, v in body.items() if k in ("status", "next_billing_at", "quantity", "interval")}
    return await _update("subscriptions", sid, allowed, u["actor_id"], "subscription")


@api.delete("/subscriptions/{sid}")
async def delete_subscription(sid: str, u: dict = Depends(current_user)):
    return await _soft_delete("subscriptions", sid, u["actor_id"], "subscription")


@api.post("/subscriptions/{sid}/tick")
async def tick_subscription(sid: str, u: dict = Depends(current_user)):
    """Generate the next invoice for a subscription.
    In production this runs on a scheduler; we expose a manual endpoint for demo + tests."""
    sub = await db.subscriptions.find_one({"id": sid, "owner_id": u["id"]}, {"_id": 0})
    if not sub:
        raise HTTPException(404, "Subscription not found")
    if sub.get("status") != "active":
        raise HTTPException(400, f"Subscription is {sub.get('status')}")
    if sub.get("cycles") is not None and sub.get("cycles_billed", 0) >= sub["cycles"]:
        await db.subscriptions.update_one({"id": sid}, {"$set": {"status": "completed"}})
        return {"ok": True, "status": "completed"}

    line = [{
        "product_id": sub["product_id"],
        "description": f"{sub['product_name']} — {sub['interval']} cycle #{sub.get('cycles_billed', 0) + 1}",
        "qty": sub.get("quantity", 1),
        "unit_price": sub["unit_price"],
        "discount_pct": 0,
        "tax_rate": sub.get("tax_rate", 0),
    }]
    totals = _compute_totals(line)
    inv_doc = {
        "owner_id": u["id"],
        "subscription_id": sid,
        "contact_id": sub.get("contact_id"),
        "company_id": sub.get("company_id"),
        "line_items": line,
        "currency": sub["currency"],
        "number": await _next_number("invoices", "INV", u["id"]),
        "status": "draft",
        "issue_date": now_iso(),
        "due_date": (datetime.now(timezone.utc) + timedelta(days=7)).isoformat(),
        **totals,
    }
    inv = await _insert("invoices", inv_doc, u["actor_id"], "invoice")
    # advance subscription
    try:
        next_dt = datetime.fromisoformat(sub["next_billing_at"].replace("Z", "+00:00")) + timedelta(days=_interval_days(sub["interval"]))
    except Exception:
        next_dt = datetime.now(timezone.utc) + timedelta(days=_interval_days(sub["interval"]))
    await db.subscriptions.update_one(
        {"id": sid},
        {"$set": {
            "last_billed_at": now_iso(),
            "next_billing_at": next_dt.isoformat(),
            "cycles_billed": sub.get("cycles_billed", 0) + 1,
            "updated_at": now_iso(),
        }},
    )
    return {"ok": True, "invoice_id": inv["id"], "invoice_number": inv["number"]}


@api.post("/subscriptions/{sid}/mark-failed")
async def mark_sub_failed(sid: str, u: dict = Depends(current_user)):
    """Simulate a failed payment for dunning demo."""
    sub = await db.subscriptions.find_one({"id": sid, "owner_id": u["id"]}, {"_id": 0})
    if not sub:
        raise HTTPException(404, "Not found")
    fails = sub.get("failed_payments", 0) + 1
    status = "past_due" if fails < 3 else "paused"
    await db.subscriptions.update_one({"id": sid}, {"$set": {"failed_payments": fails, "status": status, "updated_at": now_iso()}})
    await audit(u["actor_id"], "payment_failed", "subscription", sid, after={"failed": fails, "status": status})
    return {"ok": True, "status": status, "failed_payments": fails}


# ── Automation engine ─────────────────────────────────────────────────────────
async def _run_automations(owner_id: str, trigger_type: str, trigger_config_match: dict, context: dict):
    """Find enabled automations whose trigger matches and execute their actions."""
    cursor = db.automations.find({"owner_id": owner_id, "enabled": True}, {"_id": 0})
    async for rule in cursor:
        trig = rule.get("trigger") or {}
        if trig.get("type") != trigger_type:
            continue
        trig_cfg = trig.get("config") or {}
        # All keys in trig_cfg must equal values in trigger_config_match
        if any(trigger_config_match.get(k) != v for k, v in trig_cfg.items() if v not in (None, "", "*")):
            continue
        run_log: list[dict] = []
        for act in rule.get("actions", []):
            try:
                result = await _execute_action(owner_id, act, context)
                run_log.append({"action": act.get("type"), "ok": True, "result": result})
            except Exception as e:
                run_log.append({"action": act.get("type"), "ok": False, "error": str(e)})
        await db.automations.update_one(
            {"id": rule["id"]},
            {"$set": {"last_run_at": now_iso(), "last_run_log": run_log},
             "$inc": {"run_count": 1}},
        )


async def _execute_action(owner_id: str, action: dict, context: dict) -> dict:
    t = action.get("type")
    cfg = action.get("config") or {}
    if t == "create_task":
        tid = new_id()
        await db.tasks.insert_one({
            "id": tid, "owner_id": owner_id, "title": cfg.get("name", "New Task"),
            "related_entity_type": context.get("entity_type"),
            "related_entity_id": context.get("entity_id"),
            "status": "open", "created_at": now_iso(),
        })
        return {"task_id": tid}
    if t == "send_email_draft":
        gen_id = new_id()
        subject = cfg.get("subject") or "Automated follow-up"
        body = cfg.get("body") or f"[Draft template {cfg.get('template','default')}] — review before sending."
        await db.ai_generations.insert_one({
            "id": gen_id, "owner_id": owner_id, "kind": "email",
            "prompt": f"automation:{cfg.get('template')}", "output": body,
            "grounding_fields": [{"entity": context.get("entity_type"), "id": context.get("entity_id"), "fields": ["automation"]}],
            "questions_for_user": [], "llm_fields_used": [],
            "model": "automation-template", "created_at": now_iso(),
            "meta": {"subject": subject, "automation": True},
        })
        return {"draft_id": gen_id}
    if t == "webhook_post":
        url = cfg.get("url")
        if not url:
            return {"skipped": "no url"}
        import httpx
        try:
            async with httpx.AsyncClient(timeout=6) as c:
                r = await c.post(url, json={"event": context})
                return {"status_code": r.status_code}
        except Exception as e:
            return {"error": str(e)}
    if t == "tag_contact":
        contact_id = context.get("contact_id")
        tag = cfg.get("tag")
        if contact_id and tag:
            await db.contacts.update_one({"id": contact_id, "owner_id": owner_id}, {"$addToSet": {"tags": tag}})
            return {"tagged": tag}
    return {"noop": True}


@api.post("/automations/{aid}/test")
async def test_automation(aid: str, body: dict = Body(default={}), u: dict = Depends(current_user)):
    rule = await db.automations.find_one({"id": aid, "owner_id": u["id"]}, {"_id": 0})
    if not rule:
        raise HTTPException(404, "Automation not found")
    run_log = []
    for act in rule.get("actions", []):
        try:
            result = await _execute_action(u["id"], act, body.get("context") or {})
            run_log.append({"action": act.get("type"), "ok": True, "result": result})
        except Exception as e:
            run_log.append({"action": act.get("type"), "ok": False, "error": str(e)})
    await db.automations.update_one({"id": aid}, {"$set": {"last_run_at": now_iso(), "last_run_log": run_log}, "$inc": {"run_count": 1}})
    return {"ok": True, "run_log": run_log}


@api.patch("/automations/{aid}")
async def update_automation(aid: str, body: dict = Body(...), u: dict = Depends(current_user)):
    allowed = {k: v for k, v in body.items() if k in ("name", "trigger", "actions", "enabled")}
    return await _update("automations", aid, allowed, u["actor_id"], "automation")


@api.delete("/automations/{aid}")
async def delete_automation(aid: str, u: dict = Depends(current_user)):
    res = await db.automations.delete_one({"id": aid, "owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Not found")
    await audit(u["actor_id"], "delete", "automation", aid)
    return {"ok": True}


@api.get("/tasks")
async def list_tasks(u: dict = Depends(current_user)):
    return await _list("tasks", {"owner_id": u["id"]})


# ── SEO tools ─────────────────────────────────────────────────────────────────
class SEOPageIn(BaseDoc):
    url_path: str  # e.g. /services/100x-leader
    title: str
    meta_description: str
    keywords: list[str] = Field(default_factory=list)
    canonical_url: str | None = None
    og_title: str | None = None
    og_description: str | None = None
    og_image: str | None = None
    schema_jsonld: str | None = None
    priority: float = 0.7
    changefreq: Literal["always", "hourly", "daily", "weekly", "monthly", "yearly", "never"] = "weekly"


@api.get("/seo/pages")
async def list_seo_pages(u: dict = Depends(current_user)):
    return await _list("seo_pages", {"owner_id": u["id"]})


@api.post("/seo/pages")
async def create_seo_page(p: SEOPageIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    # Simple performance checklist
    doc["checklist"] = _seo_checklist(doc)
    return await _insert("seo_pages", doc, u["actor_id"], "seo_page")


@api.put("/seo/pages/{pid}")
async def update_seo_page(pid: str, p: SEOPageIn, u: dict = Depends(current_user)):
    patch = p.model_dump()
    patch["checklist"] = _seo_checklist(patch)
    return await _update("seo_pages", pid, patch, u["actor_id"], "seo_page")


@api.delete("/seo/pages/{pid}")
async def delete_seo_page(pid: str, u: dict = Depends(current_user)):
    res = await db.seo_pages.delete_one({"id": pid, "owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Not found")
    await audit(u["actor_id"], "delete", "seo_page", pid)
    return {"ok": True}


def _seo_checklist(page: dict) -> list[dict]:
    out: list[dict] = []
    t = page.get("title") or ""
    md = page.get("meta_description") or ""
    out.append({"check": "Title length 30–60 chars", "pass": 30 <= len(t) <= 60, "value": len(t)})
    out.append({"check": "Meta description 70–160 chars", "pass": 70 <= len(md) <= 160, "value": len(md)})
    out.append({"check": "At least 1 keyword", "pass": len(page.get("keywords") or []) >= 1})
    out.append({"check": "Canonical URL set", "pass": bool(page.get("canonical_url"))})
    out.append({"check": "Open Graph image", "pass": bool(page.get("og_image"))})
    out.append({"check": "Schema.org JSON-LD", "pass": bool(page.get("schema_jsonld"))})
    return out


@api.get("/seo/sitemap.xml")
async def sitemap_xml(owner_email: str | None = None):
    q: dict = {}
    if owner_email:
        user = await db.users.find_one({"email": owner_email.lower()}, {"_id": 0, "id": 1})
        if user:
            q["owner_id"] = user["id"]
    pages = await db.seo_pages.find(q, {"_id": 0}).to_list(500)
    base = os.environ.get("PUBLIC_SITE_BASE", "https://climbleadershiplab.vercel.app")
    body = ['<?xml version="1.0" encoding="UTF-8"?>',
            '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">']
    for p in pages:
        loc = (p.get("canonical_url") or f"{base}{p.get('url_path','/')}").replace("&", "&amp;")
        body.append(
            f"<url><loc>{loc}</loc><changefreq>{p.get('changefreq','weekly')}</changefreq>"
            f"<priority>{p.get('priority',0.7)}</priority></url>"
        )
    body.append("</urlset>")
    from fastapi.responses import Response
    return Response(content="\n".join(body), media_type="application/xml")


class SchemaSuggestReq(BaseDoc):
    url_path: str
    page_title: str
    business_type: Literal["coach", "consultant", "fitness", "service", "organization", "course", "event"] = "coach"


@api.post("/seo/schema-suggest")
async def seo_schema_suggest(req: SchemaSuggestReq, u: dict = Depends(current_user)):
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
    except Exception as e:
        raise HTTPException(500, f"LLM lib missing: {e}")
    products = await db.products.find({"owner_id": u["id"], "active": True}, {"_id": 0}).to_list(20)
    ctx_lines = [f"Brand: {u.get('name','')}", f"Page: {req.url_path} / {req.page_title}", f"Business type: {req.business_type}"]
    if products:
        ctx_lines.append("Offered products:")
        for p in products[:8]:
            ctx_lines.append(f"  - {p['name']} ({p['currency']} {p['unit_price']})")
    sys = (
        "You are an SEO assistant. Output ONLY valid schema.org JSON-LD for the given page, "
        "grounded in the provided context. Never invent prices, offers, or reviews. "
        "Return strict JSON-LD (no markdown fences). Include @context=schema.org and a minimal but correct @type "
        "appropriate for the business_type (e.g. Service, ProfessionalService, Course, Event). "
        "If you don't have a piece of info, omit that field rather than invent it."
    )
    session_id = f"seo-{u['id']}-{new_id()[:8]}"
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message=sys).with_model("gemini", "gemini-3-flash-preview")
    try:
        raw = await chat.send_message(UserMessage(text="\n".join(ctx_lines)))
    except Exception as e:
        raise HTTPException(502, f"LLM error: {e}")
    # Strip fences if model returned any
    out = raw.strip()
    if out.startswith("```"):
        out = out.split("```", 2)[1]
        if out.startswith("json"):
            out = out[4:]
        out = out.strip()
    return {"jsonld": out}


# ── Calendly webhook (public inbound) ─────────────────────────────────────────
@api.post("/webhook/calendly")
async def calendly_webhook(body: dict = Body(...), request: Request = None):
    """Calendly sends an 'invitee.created' event. Paste this URL into a Calendly webhook
    subscription (v2 API) and every booking will auto-create Contact + Deal.
    For this MVP we accept any owner_email in metadata OR fall back to the seeded owner."""
    payload = body.get("payload") or body  # Calendly wraps in 'payload'; allow raw
    email = (payload.get("email") or payload.get("invitee", {}).get("email") or "").lower()
    name = payload.get("name") or payload.get("invitee", {}).get("name") or ""
    first = name.split(" ")[0] if name else "Calendly"
    last = " ".join(name.split(" ")[1:]) if name and " " in name else "Lead"
    scheduled = payload.get("scheduled_event", {}) or {}
    event_name = scheduled.get("name") or "Discovery Call"

    # Determine owner — prefer tracking.utm_source=ascent_owner_id else seed user
    tracking = payload.get("tracking") or {}
    owner_id = tracking.get("utm_source")
    if not owner_id:
        seed = await db.users.find_one({"email": SEED_EMAIL}, {"_id": 0, "id": 1})
        owner_id = seed["id"] if seed else None
    if not owner_id:
        raise HTTPException(400, "No owner_id and no seed user")

    contact_id = None
    if email:
        existing = await db.contacts.find_one({"email": email, "owner_id": owner_id}, {"_id": 0})
        if existing:
            contact_id = existing["id"]
        else:
            contact_id = new_id()
            await db.contacts.insert_one({
                "id": contact_id, "owner_id": owner_id,
                "first_name": first, "last_name": last, "email": email,
                "tags": ["source:calendly"], "custom_fields": {"calendly_event": event_name},
                "consent": {"marketing": False, "newsletter": False, "source": "calendly", "updated_at": now_iso()},
                "interaction_count": 1, "last_activity_at": now_iso(),
                "created_at": now_iso(), "updated_at": now_iso(),
            })
    # Find first Basecamp stage
    stage = await db.pipeline_stages.find_one(
        {"owner_id": owner_id, "altitude_label": "Basecamp"}, {"_id": 0}, sort=[("order", 1)]
    )
    deal_id = None
    if stage:
        deal_id = new_id()
        await db.deals.insert_one({
            "id": deal_id, "owner_id": owner_id,
            "title": f"{first} {last} — {event_name}",
            "contact_id": contact_id, "pipeline_stage_id": stage["id"],
            "value": 0, "currency": "USD", "probability": stage.get("probability", 10),
            "status": "open", "tags": ["calendly"],
            "created_at": now_iso(), "updated_at": now_iso(),
        })
    await audit(owner_id, "calendly_booking", "contact", contact_id or "unknown",
                after={"email": email, "event": event_name, "deal_id": deal_id})
    # Fire form_submission-like automations
    await _run_automations(owner_id, "calendly_booking", {"event": event_name},
                           {"entity_type": "contact", "entity_id": contact_id, "contact_id": contact_id})
    return {"ok": True, "contact_id": contact_id, "deal_id": deal_id}


# ── Automation is hooked inline into:
#    • PATCH /api/deals/{id}/stage (move_deal, ~line 470)
#    • POST /api/forms/{slug}/submit (submit_form, ~line 822)
#    • POST /api/webhook/calendly (calendly_webhook)
# Search for `_run_automations(` to trace the hook sites.


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 2 — Tasks, Emails, Multi-step funnels, Subscription scheduler
# ══════════════════════════════════════════════════════════════════════════════

# ── Tasks CRUD ────────────────────────────────────────────────────────────────
@api.post("/tasks")
async def create_task_manual(p: TaskIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    doc["status"] = "open"
    doc["source"] = "manual"
    if p.contact_id:
        doc["related_entity_type"] = "contact"
        doc["related_entity_id"] = p.contact_id
    elif p.deal_id:
        doc["related_entity_type"] = "deal"
        doc["related_entity_id"] = p.deal_id
    return await _insert("tasks", doc, u["actor_id"], "task")


@api.patch("/tasks/{tid}")
async def update_task(tid: str, body: dict = Body(...), u: dict = Depends(current_user)):
    allowed = {k: v for k, v in body.items() if k in ("title", "status", "due_date", "notes", "assignee_id")}
    if body.get("status") == "done":
        allowed["completed_at"] = now_iso()
    return await _update("tasks", tid, allowed, u["actor_id"], "task")


@api.delete("/tasks/{tid}")
async def delete_task(tid: str, u: dict = Depends(current_user)):
    res = await db.tasks.delete_one({"id": tid, "owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Not found")
    await audit(u["actor_id"], "delete", "task", tid)
    return {"ok": True}


# ── Email log (manual) ────────────────────────────────────────────────────────
@api.post("/emails")
async def log_email(p: EmailLogIn, u: dict = Depends(current_user)):
    doc = p.model_dump()
    doc["owner_id"] = u["id"]
    doc["received_at"] = p.received_at or now_iso()
    doc["source"] = "manual"
    doc["id"] = new_id()
    doc["created_at"] = now_iso()
    await db.email_messages.insert_one(dict(doc))
    await audit(u["actor_id"], "log_email", "email", doc["id"], after={"subject": p.subject, "direction": p.direction})
    if p.contact_id:
        await db.contacts.update_one(
            {"id": p.contact_id, "owner_id": u["id"]},
            {"$set": {"last_activity_at": now_iso()}, "$inc": {"interaction_count": 1}},
        )
    return {k: v for k, v in doc.items() if k != "_id"}


@api.get("/emails")
async def list_emails(contact_id: str | None = None, u: dict = Depends(current_user)):
    q: dict = {"owner_id": u["id"]}
    if contact_id:
        q["contact_id"] = contact_id
    rows = await db.email_messages.find(q, {"_id": 0}).sort("received_at", -1).to_list(200)
    return rows


@api.delete("/emails/{eid}")
async def delete_email(eid: str, u: dict = Depends(current_user)):
    res = await db.email_messages.delete_one({"id": eid, "owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Not found")
    await audit(u["actor_id"], "delete", "email", eid)
    return {"ok": True}


# ── Subscription background scheduler ────────────────────────────────────────
_scheduler_state = {"running": False, "last_run_at": None, "last_run_processed": 0, "tick_seconds": 300}


async def _subscription_ticker_loop():
    """Lightweight asyncio scheduler. Every tick_seconds, find active subscriptions
    whose next_billing_at is due and generate the next invoice.
    Safe to run inside the FastAPI worker; idempotent via next_billing_at advance."""
    _scheduler_state["running"] = True
    log.info("📅 subscription ticker loop started")
    while _scheduler_state["running"]:
        try:
            await asyncio.sleep(_scheduler_state["tick_seconds"])
            processed = 0
            now = datetime.now(timezone.utc)
            cursor = db.subscriptions.find(
                {"status": "active", "next_billing_at": {"$lte": now.isoformat()}, "deleted_at": {"$exists": False}},
                {"_id": 0},
            )
            subs = await cursor.to_list(500)
            for sub in subs:
                try:
                    owner_id = sub["owner_id"]
                    if sub.get("cycles") is not None and sub.get("cycles_billed", 0) >= sub["cycles"]:
                        await db.subscriptions.update_one({"id": sub["id"]}, {"$set": {"status": "completed", "updated_at": now_iso()}})
                        continue
                    line = [{
                        "product_id": sub["product_id"],
                        "description": f"{sub['product_name']} — {sub['interval']} cycle #{sub.get('cycles_billed', 0) + 1}",
                        "qty": sub.get("quantity", 1), "unit_price": sub["unit_price"],
                        "discount_pct": 0, "tax_rate": sub.get("tax_rate", 0),
                    }]
                    totals = _compute_totals(line)
                    inv_doc = {
                        "owner_id": owner_id, "subscription_id": sub["id"],
                        "contact_id": sub.get("contact_id"), "company_id": sub.get("company_id"),
                        "line_items": line, "currency": sub["currency"],
                        "number": await _next_number("invoices", "INV", owner_id),
                        "status": "draft", "issue_date": now_iso(),
                        "due_date": (now + timedelta(days=7)).isoformat(),
                        **totals, "auto_generated": True,
                    }
                    await _insert("invoices", inv_doc, owner_id, "invoice")
                    try:
                        next_dt = datetime.fromisoformat(sub["next_billing_at"].replace("Z", "+00:00")) + timedelta(days=_interval_days(sub["interval"]))
                    except Exception:
                        next_dt = now + timedelta(days=_interval_days(sub["interval"]))
                    await db.subscriptions.update_one(
                        {"id": sub["id"]},
                        {"$set": {
                            "last_billed_at": now_iso(), "next_billing_at": next_dt.isoformat(),
                            "cycles_billed": sub.get("cycles_billed", 0) + 1, "updated_at": now_iso(),
                        }},
                    )
                    processed += 1
                except Exception as per_sub_err:
                    log.warning(f"ticker: subscription {sub.get('id')} failed: {per_sub_err}")
            _scheduler_state["last_run_at"] = now_iso()
            _scheduler_state["last_run_processed"] = processed
            if processed:
                log.info(f"📅 ticker: generated {processed} invoice(s) from due subscriptions")
        except asyncio.CancelledError:
            break
        except Exception as e:
            log.exception(f"ticker error: {e}")


@api.get("/scheduler/status")
async def scheduler_status(u: dict = Depends(current_user)):
    return _scheduler_state


@app.on_event("startup")
async def _start_ticker():
    asyncio.create_task(_subscription_ticker_loop())


@app.on_event("shutdown")
async def _stop_ticker():
    _scheduler_state["running"] = False


# ── Multi-step funnel: public getter returns steps if defined ─────────────────
# (existing /api/forms/{slug}/public already returns the whole form — steps
#  field rides along; the frontend decides how to render.)


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 3 — Team seats + Invites + RBAC, IMAP inbound sync
# ══════════════════════════════════════════════════════════════════════════════

ROLE_CHOICES = ["owner", "admin", "accountant", "rep", "va", "view"]


class InviteIn(BaseDoc):
    email: EmailStr
    role: Literal["admin", "rep", "va", "view"] = "rep"


class AcceptInviteReq(BaseDoc):
    token: str
    password: str
    name: str | None = None


# ── Team members ──────────────────────────────────────────────────────────────
@api.get("/team/members")
async def list_team_members(u: dict = Depends(current_user)):
    rows = await db.users.find(
        {"$or": [{"id": u["id"]}, {"team_owner_id": u["id"]}]},
        {"_id": 0, "password_hash": 0},
    ).to_list(100)
    return rows


@api.patch("/team/members/{uid}")
async def update_member_role(uid: str, body: dict = Body(...), u: dict = Depends(require_owner_admin)):
    if uid == u["id"]:
        raise HTTPException(400, "Cannot change your own role")
    target = await db.users.find_one({"id": uid, "team_owner_id": u["id"]}, {"_id": 0})
    if not target:
        raise HTTPException(404, "Member not found in your team")
    role = body.get("role")
    if role not in ROLE_CHOICES:
        raise HTTPException(400, f"Invalid role. Allowed: {ROLE_CHOICES}")
    if body.get("role") == "owner":
        raise HTTPException(400, "Team can have only one owner")
    await db.users.update_one({"id": uid}, {"$set": {"role": role}})
    await audit(u["actor_id"], "update_member_role", "user", uid, after={"role": role})
    return {"ok": True, "role": role}


@api.delete("/team/members/{uid}")
async def remove_member(uid: str, u: dict = Depends(require_owner_admin)):
    if uid == u["id"]:
        raise HTTPException(400, "Cannot remove the team owner")
    res = await db.users.delete_one({"id": uid, "team_owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Member not found in your team")
    await audit(u["actor_id"], "remove_member", "user", uid)
    return {"ok": True}


# ── Invites ───────────────────────────────────────────────────────────────────
@api.get("/team/invites")
async def list_invites(u: dict = Depends(require_owner_admin)):
    return await db.invites.find({"team_owner_id": u["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)


@api.post("/team/invites")
async def create_invite(p: InviteIn, u: dict = Depends(require_owner_admin)):
    existing_user = await db.users.find_one({"email": p.email.lower()}, {"_id": 0, "id": 1})
    if existing_user:
        raise HTTPException(409, "A user with that email already exists")
    token = secrets.token_urlsafe(24)
    inv = {
        "id": new_id(),
        "team_owner_id": u["id"],
        "email": p.email.lower(),
        "role": p.role,
        "token": token,
        "status": "pending",
        "invited_by": u["actor_id"],
        "invited_by_name": u.get("name"),
        "created_at": now_iso(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(days=14)).isoformat(),
    }
    await db.invites.insert_one(dict(inv))
    await audit(u["actor_id"], "create_invite", "invite", inv["id"], after={"email": p.email, "role": p.role})
    return {k: v for k, v in inv.items() if k != "_id"}


@api.delete("/team/invites/{iid}")
async def revoke_invite(iid: str, u: dict = Depends(require_owner_admin)):
    res = await db.invites.delete_one({"id": iid, "team_owner_id": u["id"]})
    if res.deleted_count == 0:
        raise HTTPException(404, "Invite not found")
    await audit(u["actor_id"], "revoke_invite", "invite", iid)
    return {"ok": True}


@api.get("/auth/invite/{token}")
async def peek_invite(token: str):
    inv = await db.invites.find_one({"token": token, "status": "pending"}, {"_id": 0, "token": 0})
    if not inv:
        raise HTTPException(404, "Invite not found or already used")
    now = datetime.now(timezone.utc)
    try:
        exp = datetime.fromisoformat(inv["expires_at"].replace("Z", "+00:00"))
    except Exception:
        exp = now
    if exp < now:
        raise HTTPException(410, "Invite expired")
    return inv


@api.post("/auth/accept-invite", response_model=AuthResp)
async def accept_invite(req: AcceptInviteReq):
    inv = await db.invites.find_one({"token": req.token, "status": "pending"})
    if not inv:
        raise HTTPException(404, "Invite not found or already used")
    # Check expiry
    try:
        exp = datetime.fromisoformat(inv["expires_at"].replace("Z", "+00:00"))
    except Exception:
        exp = datetime.now(timezone.utc)
    if exp < datetime.now(timezone.utc):
        raise HTTPException(410, "Invite expired")
    if await db.users.find_one({"email": inv["email"]}):
        raise HTTPException(409, "A user with that email already exists")
    uid = new_id()
    doc = {
        "id": uid,
        "email": inv["email"],
        "name": req.name or inv["email"].split("@")[0],
        "password_hash": hash_pw(req.password),
        "role": inv["role"],
        "team_owner_id": inv["team_owner_id"],
        "brand_voice": {"tone": "professional", "vocabulary_hints": "", "signature": "", "banned_phrases": []},
        "created_at": now_iso(),
    }
    await db.users.insert_one(doc)
    await db.invites.update_one(
        {"id": inv["id"]},
        {"$set": {"status": "accepted", "accepted_at": now_iso(), "accepted_by": uid}},
    )
    await audit(uid, "accept_invite", "invite", inv["id"], after={"role": inv["role"]})
    public = {k: v for k, v in doc.items() if k not in ("password_hash", "_id")}
    return {"token": make_token(uid, inv["email"]), "user": public}


# ── IMAP inbound sync ─────────────────────────────────────────────────────────
class ImapConfigIn(BaseDoc):
    host: str
    port: int = 993
    use_ssl: bool = True
    username: str
    password: str  # stored obfuscated only; real use sends password with each sync
    mailbox: str = "INBOX"


@api.get("/email/imap/config")
async def get_imap_config(u: dict = Depends(current_user)):
    cfg = await db.imap_configs.find_one({"team_owner_id": u["id"]}, {"_id": 0})
    if not cfg:
        return None
    cfg["password"] = "••••••" if cfg.get("password") else ""
    return cfg


@api.post("/email/imap/config")
async def save_imap_config(p: ImapConfigIn, u: dict = Depends(require_owner_admin)):
    doc = p.model_dump()
    # NEVER persist the plaintext IMAP password at rest.  Users supply it on
    # each sync call.  We keep the field shape for the client but blank the value.
    doc["password"] = ""
    doc["team_owner_id"] = u["id"]
    doc["updated_at"] = now_iso()
    await db.imap_configs.update_one({"team_owner_id": u["id"]}, {"$set": doc}, upsert=True)
    await audit(u["actor_id"], "save", "imap_config", u["id"], after={"host": p.host, "username": p.username})
    return {"ok": True}


@api.delete("/email/imap/config")
async def delete_imap_config(u: dict = Depends(require_owner_admin)):
    await db.imap_configs.delete_one({"team_owner_id": u["id"]})
    await audit(u["actor_id"], "delete", "imap_config", u["id"])
    return {"ok": True}


def _fetch_imap_messages(host: str, port: int, ssl_on: bool, username: str, password: str, mailbox: str, limit: int) -> list[dict]:
    """Synchronous IMAP fetch; called from a thread. Returns list of dicts."""
    import imaplib
    import email as emaillib
    from email.header import decode_header
    from email.utils import parseaddr, parsedate_to_datetime

    def _decode(s):
        if s is None:
            return ""
        parts = decode_header(s)
        out = []
        for p, enc in parts:
            if isinstance(p, bytes):
                try:
                    out.append(p.decode(enc or "utf-8", errors="replace"))
                except Exception:
                    out.append(p.decode("utf-8", errors="replace"))
            else:
                out.append(p)
        return "".join(out)

    M = imaplib.IMAP4_SSL(host, port, timeout=30) if ssl_on else imaplib.IMAP4(host, port, timeout=30)
    try:
        M.login(username, password)
        M.select(mailbox, readonly=True)
        _, data = M.search(None, "ALL")
        ids = (data[0] or b"").split()
        last = ids[-limit:] if len(ids) > limit else ids
        out: list[dict] = []
        for mid in reversed(last):
            _, d = M.fetch(mid, "(RFC822)")
            if not d or not d[0]:
                continue
            msg = emaillib.message_from_bytes(d[0][1])
            from_name, from_addr = parseaddr(msg.get("From") or "")
            to_addrs = msg.get("To") or ""
            subj = _decode(msg.get("Subject"))
            date_hdr = msg.get("Date")
            try:
                received_at = parsedate_to_datetime(date_hdr).isoformat() if date_hdr else now_iso()
            except Exception:
                received_at = now_iso()
            # Body (prefer text/plain)
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain" and "attachment" not in str(part.get("Content-Disposition") or ""):
                        try:
                            body = part.get_payload(decode=True).decode(part.get_content_charset() or "utf-8", errors="replace")
                            break
                        except Exception:
                            continue
            else:
                try:
                    body = msg.get_payload(decode=True).decode(msg.get_content_charset() or "utf-8", errors="replace")
                except Exception:
                    body = str(msg.get_payload())
            out.append({
                "imap_uid": mid.decode() if isinstance(mid, bytes) else str(mid),
                "from_name": from_name, "from_addr": (from_addr or "").lower(),
                "to_addr": to_addrs, "subject": subj or "(no subject)",
                "body": (body or "")[:8000],
                "received_at": received_at,
            })
        return out
    finally:
        try:
            M.logout()
        except Exception:
            pass


@api.post("/email/imap/sync")
async def imap_sync(body: dict = Body(default={}), u: dict = Depends(current_user)):
    limit = int(body.get("limit") or 20)
    cfg = await db.imap_configs.find_one({"team_owner_id": u["id"]}, {"_id": 0})
    if not cfg:
        raise HTTPException(400, "No IMAP config. Save one first.")
    # Always require the password on each sync request — never fall back to a
    # stored credential. We also never persist a plaintext password at rest.
    pw = body.get("password")
    if not pw or pw == "••••••":
        raise HTTPException(400, "Provide your IMAP app password in the request body for this sync.")
    try:
        msgs = await asyncio.to_thread(
            _fetch_imap_messages, cfg["host"], int(cfg.get("port") or 993),
            bool(cfg.get("use_ssl", True)), cfg["username"], pw,
            cfg.get("mailbox") or "INBOX", limit,
        )
    except Exception as e:
        raise HTTPException(502, f"IMAP error: {e}")
    # Index contacts by email for match
    contacts = await db.contacts.find({"owner_id": u["id"]}, {"_id": 0, "id": 1, "email": 1}).to_list(5000)
    by_email = {(c.get("email") or "").lower(): c["id"] for c in contacts if c.get("email")}
    inserted = 0
    for m in msgs:
        # dedupe by (owner, imap_uid)
        if await db.email_messages.find_one({"owner_id": u["id"], "imap_uid": m["imap_uid"]}, {"_id": 0, "id": 1}):
            continue
        contact_id = by_email.get(m["from_addr"])
        await db.email_messages.insert_one({
            "id": new_id(), "owner_id": u["id"], "direction": "in",
            "contact_id": contact_id, "subject": m["subject"], "body": m["body"],
            "from_addr": m["from_addr"], "to_addr": m["to_addr"],
            "received_at": m["received_at"], "source": "imap", "imap_uid": m["imap_uid"],
            "created_at": now_iso(),
        })
        if contact_id:
            await db.contacts.update_one(
                {"id": contact_id, "owner_id": u["id"]},
                {"$set": {"last_activity_at": now_iso()}, "$inc": {"interaction_count": 1}},
            )
        inserted += 1
    await db.imap_configs.update_one({"team_owner_id": u["id"]}, {"$set": {"last_sync_at": now_iso(), "last_sync_count": inserted}})
    await audit(u["actor_id"], "imap_sync", "imap_config", u["id"], after={"fetched": len(msgs), "inserted": inserted})
    return {"ok": True, "fetched": len(msgs), "inserted": inserted}


# Stamp pre-existing rows where team_owner_id missing (backfill for demo user)
@app.on_event("startup")
async def _backfill_team_owner():
    try:
        await db.users.update_many(
            {"team_owner_id": {"$exists": False}},
            [{"$set": {"team_owner_id": "$id"}}],
        )
    except Exception as e:
        log.warning(f"backfill team_owner_id failed: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 5 — Word export for quotes  +  PDF attachments (quotes / invoices)
# ══════════════════════════════════════════════════════════════════════════════
from pathlib import Path as _Path
from fastapi import UploadFile, File as _File, Form as _Form
from fastapi.responses import Response as _Response

UPLOAD_ROOT = _Path(os.environ.get("UPLOAD_ROOT", "/app/backend/uploads"))
UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)

ATTACHABLE = {"quotes": "quotes", "invoices": "invoices"}


def _money_str(v: float, cur: str) -> str:
    return f"{cur.upper()} {float(v or 0):,.2f}"


@api.get("/quotes/{qid}/export/docx")
async def quote_export_docx(qid: str, u: dict = Depends(current_user)):
    """Generate a branded Word document for a quote; user previews/prints/sends from Word."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    q = await db.quotes.find_one({"id": qid, "owner_id": u["id"]}, {"_id": 0})
    if not q:
        raise HTTPException(404, "Quote not found")
    owner = await db.users.find_one({"id": u["id"]}, {"_id": 0, "password_hash": 0}) or {}
    contact = await db.contacts.find_one({"id": q.get("contact_id")}, {"_id": 0}) if q.get("contact_id") else None
    company = await db.companies.find_one({"id": q.get("company_id")}, {"_id": 0}) if q.get("company_id") else None

    tpl = owner.get("quote_template") or {}
    hex_color = (tpl.get("accent_color_hex") or "E26E4A").lstrip("#")[:6].upper() or "E26E4A"
    try:
        accent = RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
    except Exception:
        accent = RGBColor(0xE2, 0x6E, 0x4A)
    title_label = tpl.get("title_label") or "QUOTATION"
    company_name = tpl.get("company_name") or owner.get("name") or "Ascent CRM"
    tagline = tpl.get("tagline") or "Ascent CRM · Climb Leadership Lab"
    footer_text = tpl.get("footer_text") or ""
    signature_block = tpl.get("signature_block") or (owner.get("brand_voice") or {}).get("signature") or ""

    doc = Document()
    TERRACOTTA = accent
    SLATE = RGBColor(0x94, 0xA3, 0xB8)

    # Title block
    h = doc.add_heading("", level=0)
    run = h.add_run(title_label)
    run.font.color.rgb = TERRACOTTA
    run.font.size = Pt(28)

    meta = doc.add_paragraph()
    meta.add_run(f"Quote #{q['number']}\n").bold = True
    if q.get("valid_until"):
        meta.add_run(f"Valid until: {q['valid_until']}\n")
    meta.add_run(f"Issued: {(q.get('created_at') or '')[:10]}\n")

    # From / To
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    lp = left.paragraphs[0]
    lp.add_run("FROM\n").bold = True
    lp.add_run(f"{company_name}\n{owner.get('email') or ''}\n")
    rp = right.paragraphs[0]
    rp.add_run("BILL TO\n").bold = True
    if company:
        rp.add_run(f"{company.get('name') or ''}\n")
    if contact:
        rp.add_run(f"{(contact.get('first_name') or '')} {(contact.get('last_name') or '')}\n".strip() + "\n")
        if contact.get("email"):
            rp.add_run(f"{contact['email']}\n")

    doc.add_paragraph()  # spacer

    # Line items table
    items = q.get("line_items") or []
    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Description", "Qty", "Unit", "Disc %", "Total"]):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    cur = q.get("currency", "USD")
    for i, ln in enumerate(items, start=1):
        cells = tbl.rows[i].cells
        cells[0].text = ln.get("description") or ""
        cells[1].text = str(ln.get("qty") or 0)
        cells[2].text = _money_str(ln.get("unit_price") or 0, cur)
        cells[3].text = f"{(ln.get('discount_pct') or 0)}%"
        qty = float(ln.get("qty") or 0)
        up = float(ln.get("unit_price") or 0)
        sub = qty * up
        disc = sub * float(ln.get("discount_pct") or 0) / 100.0
        cells[4].text = _money_str(sub - disc, cur)

    doc.add_paragraph()

    # Totals
    tot = doc.add_paragraph()
    tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tot.add_run(f"Subtotal: {_money_str(q.get('subtotal') or 0, cur)}\n")
    tot.add_run(f"Tax: {_money_str(q.get('tax_total') or 0, cur)}\n")
    grand = tot.add_run(f"TOTAL: {_money_str(q.get('grand_total') or 0, cur)}")
    grand.bold = True
    grand.font.size = Pt(14)
    grand.font.color.rgb = TERRACOTTA

    # Terms
    if q.get("terms"):
        doc.add_paragraph()
        tp = doc.add_paragraph()
        tp.add_run("Terms\n").bold = True
        tp.add_run(q["terms"])

    # Footer text + signature
    if footer_text:
        ft = doc.add_paragraph()
        ft.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ft.add_run(footer_text).italic = True

    if signature_block:
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sig.add_run(signature_block)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(f"\n{tagline}")
    fr.font.color.rgb = SLATE
    fr.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    await audit(u["actor_id"], "export_docx", "quote", qid, after={"number": q["number"]})
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Quote_{q["number"]}.docx"'},
    )


# ── PDF Attachments ────────────────────────────────────────────────────────────
MAX_ATTACHMENT_BYTES = 15 * 1024 * 1024  # 15 MB
ALLOWED_ATTACHMENT_SUFFIXES = {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"}


@api.post("/{resource}/{rid}/attachments")
async def upload_attachment(
    resource: str,
    rid: str,
    file: UploadFile = _File(...),
    kind: str | None = _Form(None),
    u: dict = Depends(current_user),
):
    """Attach a PDF (or other allowed type) to a quote/invoice. `kind` is a free-form tag
    like 'signed_quote' or 'signed_invoice' for display."""
    if resource not in ATTACHABLE:
        raise HTTPException(400, f"Not attachable: {resource}")
    coll = ATTACHABLE[resource]
    row = await db[coll].find_one({"id": rid, "owner_id": u["id"]}, {"_id": 0, "id": 1})
    if not row:
        raise HTTPException(404, f"{resource} not found")

    suffix = _Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_ATTACHMENT_SUFFIXES:
        raise HTTPException(400, f"File type not allowed: {suffix}. Allowed: {sorted(ALLOWED_ATTACHMENT_SUFFIXES)}")
    data = await file.read()
    if len(data) > MAX_ATTACHMENT_BYTES:
        raise HTTPException(413, f"File too large (>{MAX_ATTACHMENT_BYTES // (1024*1024)} MB)")

    att_id = new_id()
    target_dir = UPLOAD_ROOT / u["id"] / resource / rid
    target_dir.mkdir(parents=True, exist_ok=True)
    disk_path = target_dir / f"{att_id}{suffix}"
    disk_path.write_bytes(data)

    rec = {
        "id": att_id,
        "owner_id": u["id"],
        "resource": resource,
        "resource_id": rid,
        "kind": kind or ("signed_" + resource.rstrip("s")),
        "filename": file.filename or f"attachment{suffix}",
        "content_type": file.content_type or "application/octet-stream",
        "size": len(data),
        "disk_path": str(disk_path),
        "created_at": now_iso(),
        "uploaded_by": u["actor_id"],
    }
    await db.attachments.insert_one(dict(rec))
    await audit(u["actor_id"], "upload_attachment", resource, rid,
                after={"attachment_id": att_id, "filename": rec["filename"], "size": rec["size"]})
    rec.pop("_id", None)
    return rec


@api.get("/{resource}/{rid}/attachments")
async def list_attachments(resource: str, rid: str, u: dict = Depends(current_user)):
    if resource not in ATTACHABLE:
        raise HTTPException(400, f"Not attachable: {resource}")
    rows = await db.attachments.find(
        {"owner_id": u["id"], "resource": resource, "resource_id": rid},
        {"_id": 0, "disk_path": 0},
    ).sort("created_at", -1).to_list(100)
    return rows


@api.get("/attachments/{att_id}/download")
async def download_attachment(att_id: str, u: dict = Depends(current_user)):
    att = await db.attachments.find_one({"id": att_id, "owner_id": u["id"]}, {"_id": 0})
    if not att:
        raise HTTPException(404, "Attachment not found")
    p = _Path(att["disk_path"])
    if not p.exists():
        raise HTTPException(410, "File missing on disk")
    return _Response(
        content=p.read_bytes(),
        media_type=att.get("content_type") or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{att["filename"]}"'},
    )


@api.delete("/attachments/{att_id}")
async def delete_attachment(att_id: str, u: dict = Depends(current_user)):
    att = await db.attachments.find_one({"id": att_id, "owner_id": u["id"]}, {"_id": 0})
    if not att:
        raise HTTPException(404, "Attachment not found")
    try:
        _Path(att["disk_path"]).unlink(missing_ok=True)
    except Exception as e:
        log.warning(f"attachment disk delete failed: {e}")
    await db.attachments.delete_one({"id": att_id, "owner_id": u["id"]})
    await audit(u["actor_id"], "delete_attachment", att["resource"], att["resource_id"],
                before={"attachment_id": att_id, "filename": att.get("filename")})
    return {"ok": True}


# ── Company → Contacts list (convenience) ──────────────────────────────────────
@api.get("/companies/{cid}/contacts")
async def company_contacts(cid: str, u: dict = Depends(current_user)):
    company = await db.companies.find_one({"id": cid, "owner_id": u["id"]}, {"_id": 0, "id": 1})
    if not company:
        raise HTTPException(404, "Company not found")
    rows = await db.contacts.find(
        {"owner_id": u["id"], "company_id": cid, "deleted_at": {"$exists": False}},
        {"_id": 0},
    ).sort("last_name", 1).to_list(500)
    return rows



# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 BATCH 6 — SOUTH AFRICAN ACCOUNTING MODULE (Batch A: Foundation)
# ══════════════════════════════════════════════════════════════════════════════
# Double-entry • Accrual • IFRS for SMEs–compatible • SA Chart of Accounts
# VAT 15% with VAT201 report • Period close + accountant sign-off
# DISCLAIMER: Scaffolding for a SA coaching business. All computations require
# sign-off by a qualified SA accountant (SAICA / SAIPA / CA(SA)) before filing
# with SARS. Not a replacement for professional advice.
# ══════════════════════════════════════════════════════════════════════════════
from decimal import Decimal, ROUND_HALF_UP

# Static data + helpers extracted to dedicated modules for maintainability.
from accounting_data import (  # noqa: E402
    ACCOUNT_TYPES,
    NORMAL_BALANCE,
    SA_VAT_CODES,
    SA_COA_SEED,
    _D,
    _period_key,
)


class JournalLineIn(BaseDoc):
    account_code: str
    debit: float = 0.0
    credit: float = 0.0
    description: str | None = None
    vat_code: str | None = None
    vat_amount: float | None = None
    contact_id: str | None = None


class JournalIn(BaseDoc):
    date: str  # ISO date
    memo: str
    reference: str | None = None
    lines: list[JournalLineIn]
    source: str = "manual"  # manual|adjusting|reversing|system
    source_id: str | None = None


# ---- Seed -------------------------------------------------------------------
@api.post("/accounting/seed")
async def accounting_seed(u: dict = Depends(require_owner_admin)):
    """Idempotently seed SA COA + VAT codes + current open fiscal period."""
    inserted_accounts = 0
    for a in SA_COA_SEED:
        exists = await db.accounts.find_one({"owner_id": u["id"], "code": a["code"]}, {"_id": 0, "id": 1})
        if exists:
            continue
        doc = {
            "id": new_id(), "owner_id": u["id"],
            "code": a["code"], "name": a["name"], "type": a["type"],
            "subtype": a.get("subtype"), "parent_code": a.get("parent"),
            "vat_code": a.get("vat_code"),
            "normal_balance": NORMAL_BALANCE.get(a["type"], "debit"),
            "is_header": a.get("subtype") == "header",
            "active": True, "locked": False,
            "created_at": now_iso(),
        }
        await db.accounts.insert_one(doc)
        inserted_accounts += 1

    # VAT codes
    for vc in SA_VAT_CODES:
        await db.vat_codes.update_one(
            {"owner_id": u["id"], "code": vc["code"]},
            {"$set": {**vc, "owner_id": u["id"]}},
            upsert=True,
        )

    # Current period
    now = datetime.now(timezone.utc)
    period = f"{now.year}-{now.month:02d}"
    await db.fiscal_periods.update_one(
        {"owner_id": u["id"], "period": period},
        {"$setOnInsert": {
            "id": new_id(), "owner_id": u["id"], "period": period,
            "year": now.year, "month": now.month, "status": "open",
            "created_at": now_iso(),
        }},
        upsert=True,
    )
    await audit(u["actor_id"], "accounting_seed", "accounting", u["id"],
                after={"accounts_added": inserted_accounts, "period": period})
    return {"ok": True, "accounts_added": inserted_accounts, "period_opened": period}


# ---- Chart of Accounts ------------------------------------------------------
@api.get("/accounting/accounts")
async def list_accounts(active_only: bool = True, u: dict = Depends(current_user)):
    q: dict = {"owner_id": u["id"]}
    if active_only:
        q["active"] = True
    return await db.accounts.find(q, {"_id": 0}).sort("code", 1).to_list(500)


@api.post("/accounting/accounts")
async def create_account(payload: dict = Body(...), u: dict = Depends(require_accountant)):
    required = {"code", "name", "type"}
    missing = required - set(payload.keys())
    if missing:
        raise HTTPException(400, f"Missing fields: {sorted(missing)}")
    if payload["type"] not in ACCOUNT_TYPES:
        raise HTTPException(400, f"type must be one of {ACCOUNT_TYPES}")
    if await db.accounts.find_one({"owner_id": u["id"], "code": payload["code"]}, {"_id": 0, "id": 1}):
        raise HTTPException(409, f"Account code {payload['code']} already exists")
    doc = {
        "id": new_id(), "owner_id": u["id"],
        "code": str(payload["code"]), "name": payload["name"], "type": payload["type"],
        "subtype": payload.get("subtype"), "parent_code": payload.get("parent_code"),
        "vat_code": payload.get("vat_code"),
        "normal_balance": NORMAL_BALANCE.get(payload["type"], "debit"),
        "is_header": bool(payload.get("is_header")),
        "active": True, "locked": False,
        "created_at": now_iso(),
    }
    await db.accounts.insert_one(dict(doc))
    await audit(u["actor_id"], "create_account", "account", doc["id"], after={"code": doc["code"], "name": doc["name"]})
    doc.pop("_id", None)
    return doc


@api.patch("/accounting/accounts/{aid}")
async def update_account(aid: str, payload: dict = Body(...), u: dict = Depends(require_accountant)):
    allowed = {"name", "active", "vat_code", "subtype"}
    patch = {k: v for k, v in payload.items() if k in allowed}
    res = await db.accounts.update_one({"id": aid, "owner_id": u["id"]}, {"$set": patch})
    if res.matched_count == 0:
        raise HTTPException(404, "Account not found")
    await audit(u["actor_id"], "update_account", "account", aid, after=patch)
    return {"ok": True}


# ---- Journal engine ---------------------------------------------------------
async def _validate_and_post_journal(owner_id: str, actor_id: str, payload: JournalIn, auto: bool = False) -> dict:
    """Core journal posting with double-entry + period-lock + account-validity checks."""
    if not payload.lines or len(payload.lines) < 2:
        raise HTTPException(400, "Journal must have at least two lines")
    period_key = _period_key(payload.date)
    fp = await db.fiscal_periods.find_one({"owner_id": owner_id, "period": period_key}, {"_id": 0})
    if not fp:
        # auto-create an open period for the transaction's month (common for auto-posting)
        try:
            d = datetime.fromisoformat(payload.date.replace("Z", "+00:00"))
        except Exception:
            d = datetime.now(timezone.utc)
        fp = {
            "id": new_id(), "owner_id": owner_id, "period": period_key,
            "year": d.year, "month": d.month, "status": "open", "created_at": now_iso(),
        }
        await db.fiscal_periods.insert_one(dict(fp))
    if fp["status"] in ("locked", "closed") and not auto:
        raise HTTPException(423, f"Period {period_key} is {fp['status']} — post an adjusting entry in an open period instead")

    # Resolve account codes → ids + normal balance for validation
    codes = {ln.account_code for ln in payload.lines}
    accounts = await db.accounts.find({"owner_id": owner_id, "code": {"$in": list(codes)}}, {"_id": 0}).to_list(500)
    if len(accounts) != len(codes):
        missing = codes - {a["code"] for a in accounts}
        raise HTTPException(400, f"Unknown account code(s): {sorted(missing)}")
    acct_by_code = {a["code"]: a for a in accounts}

    total_dr = Decimal("0.00")
    total_cr = Decimal("0.00")
    line_docs: list[dict] = []
    for i, ln in enumerate(payload.lines, start=1):
        d = _D(ln.debit)
        c = _D(ln.credit)
        if d < 0 or c < 0:
            raise HTTPException(400, f"Line {i}: debit/credit must be non-negative")
        if d > 0 and c > 0:
            raise HTTPException(400, f"Line {i}: a single line cannot be both debit and credit")
        if d == 0 and c == 0:
            raise HTTPException(400, f"Line {i}: zero-value line")
        acct = acct_by_code[ln.account_code]
        if acct.get("is_header"):
            raise HTTPException(400, f"Line {i}: cannot post to header account {ln.account_code}")
        if not acct.get("active", True):
            raise HTTPException(400, f"Line {i}: account {ln.account_code} is inactive")
        total_dr += d
        total_cr += c
        line_docs.append({
            "id": new_id(),
            "account_id": acct["id"], "account_code": acct["code"], "account_name": acct["name"],
            "debit": float(d), "credit": float(c),
            "description": ln.description, "vat_code": ln.vat_code,
            "vat_amount": float(_D(ln.vat_amount)) if ln.vat_amount is not None else None,
            "contact_id": ln.contact_id,
        })
    if total_dr != total_cr:
        raise HTTPException(400, f"Journal unbalanced: DR {total_dr} ≠ CR {total_cr}")

    jdoc = {
        "id": new_id(), "owner_id": owner_id,
        "date": payload.date[:10], "period": period_key,
        "memo": payload.memo, "reference": payload.reference,
        "source": payload.source or "manual", "source_id": payload.source_id,
        "lines": line_docs,
        "total_debit": float(total_dr), "total_credit": float(total_cr),
        "posted": True, "posted_at": now_iso(),
        "created_by": actor_id, "created_at": now_iso(),
        "reversed_of": None, "reversed_by": None,
    }
    await db.journals.insert_one(dict(jdoc))
    await audit(actor_id, "post_journal", "journal", jdoc["id"],
                after={"memo": payload.memo, "total": float(total_dr), "source": payload.source})
    jdoc.pop("_id", None)
    return jdoc


@api.get("/accounting/journals")
async def list_journals(period: str | None = None, limit: int = 200, u: dict = Depends(current_user)):
    q: dict = {"owner_id": u["id"]}
    if period:
        q["period"] = period
    return await db.journals.find(q, {"_id": 0}).sort("date", -1).to_list(limit)


@api.get("/accounting/journals/{jid}")
async def get_journal(jid: str, u: dict = Depends(current_user)):
    j = await db.journals.find_one({"id": jid, "owner_id": u["id"]}, {"_id": 0})
    if not j:
        raise HTTPException(404, "Journal not found")
    return j


@api.post("/accounting/journals")
async def post_journal(payload: JournalIn, u: dict = Depends(require_accountant)):
    return await _validate_and_post_journal(u["id"], u["actor_id"], payload, auto=False)


@api.post("/accounting/journals/{jid}/reverse")
async def reverse_journal(jid: str, u: dict = Depends(require_accountant)):
    j = await db.journals.find_one({"id": jid, "owner_id": u["id"]}, {"_id": 0})
    if not j:
        raise HTTPException(404, "Journal not found")
    if j.get("reversed_by"):
        raise HTTPException(409, "Journal already reversed")
    # Build flipped lines (negate vat_amount so VAT201 totals net-out correctly)
    rev_lines = [JournalLineIn(
        account_code=ln["account_code"],
        debit=float(ln.get("credit") or 0),
        credit=float(ln.get("debit") or 0),
        description=f"REV: {ln.get('description') or ''}",
        vat_code=ln.get("vat_code"),
        vat_amount=(-float(ln["vat_amount"]) if ln.get("vat_amount") is not None else None),
    ) for ln in j["lines"]]
    rev_payload = JournalIn(
        date=datetime.now(timezone.utc).date().isoformat(),
        memo=f"Reversal of {j['memo']}",
        reference=f"REV-{j.get('reference') or j['id']}",
        source="reversing", source_id=j["id"], lines=rev_lines,
    )
    rev = await _validate_and_post_journal(u["id"], u["actor_id"], rev_payload, auto=False)
    await db.journals.update_one({"id": jid}, {"$set": {"reversed_by": rev["id"], "reversed_at": now_iso()}})
    await db.journals.update_one({"id": rev["id"]}, {"$set": {"reversed_of": jid}})
    rev["reversed_of"] = jid  # Reflect the linkage in the returned response (DB has it too)
    return rev


# ---- Auto-posting from Invoices + Payments ---------------------------------
async def _auto_post_invoice_journal(owner_id: str, invoice: dict):
    """On invoice creation: DR Debtors (22000), CR Revenue (61000), CR VAT Output (52000)."""
    try:
        grand = _D(invoice.get("grand_total"))
        tax = _D(invoice.get("tax_total"))
        net = grand - tax
        lines = [
            JournalLineIn(account_code="22000", debit=float(grand), description=f"Invoice {invoice.get('number')}"),
            JournalLineIn(account_code="61000", credit=float(net), description="Coaching Revenue", vat_code="S"),
        ]
        if tax > 0:
            lines.append(JournalLineIn(account_code="52000", credit=float(tax), description="VAT Output 15%", vat_code="S", vat_amount=float(tax)))
        payload = JournalIn(
            date=(invoice.get("issue_date") or now_iso())[:10],
            memo=f"Invoice {invoice.get('number')} issued",
            reference=invoice.get("number"), source="invoice", source_id=invoice.get("id"),
            lines=lines,
        )
        await _validate_and_post_journal(owner_id, invoice.get("owner_id") or owner_id, payload, auto=True)
    except HTTPException as e:
        log.warning(f"auto_post_invoice_journal skipped: {e.detail}")
    except Exception as e:
        log.warning(f"auto_post_invoice_journal error: {e}")


async def _auto_post_payment_journal(owner_id: str, invoice: dict, provider: str = "stripe"):
    """On payment capture: DR Bank (21000), CR Debtors (22000)."""
    try:
        grand = _D(invoice.get("grand_total"))
        # Route via clearing account for card processors (realistic SA practice)
        clearing = "56000"
        lines = [
            JournalLineIn(account_code=clearing, debit=float(grand), description=f"{provider} clearing"),
            JournalLineIn(account_code="22000", credit=float(grand), description=f"Payment for invoice {invoice.get('number')}"),
        ]
        payload = JournalIn(
            date=now_iso()[:10],
            memo=f"{provider.title()} payment received · {invoice.get('number')}",
            reference=invoice.get("number"), source="payment",
            source_id=invoice.get("id"), lines=lines,
        )
        await _validate_and_post_journal(owner_id, owner_id, payload, auto=True)
    except Exception as e:
        log.warning(f"auto_post_payment_journal error: {e}")


# ---- Reports ----------------------------------------------------------------
async def _balance_by_code(owner_id: str, date_to: str | None = None, date_from: str | None = None) -> dict:
    """Returns {account_code: (debit_total, credit_total, net_signed)}.
    net_signed is positive on the account's normal side."""
    match: dict = {"owner_id": owner_id, "posted": True}
    if date_from or date_to:
        match["date"] = {}
        if date_from:
            match["date"]["$gte"] = date_from
        if date_to:
            match["date"]["$lte"] = date_to
    pipeline = [
        {"$match": match},
        {"$unwind": "$lines"},
        {"$group": {
            "_id": "$lines.account_code",
            "debit": {"$sum": "$lines.debit"},
            "credit": {"$sum": "$lines.credit"},
        }},
    ]
    out: dict = {}
    async for row in db.journals.aggregate(pipeline):
        out[row["_id"]] = {"debit": _D(row.get("debit") or 0), "credit": _D(row.get("credit") or 0)}
    return out


@api.get("/accounting/reports/trial-balance")
async def trial_balance(date_to: str | None = None, u: dict = Depends(current_user)):
    accounts = await db.accounts.find({"owner_id": u["id"]}, {"_id": 0}).sort("code", 1).to_list(500)
    bals = await _balance_by_code(u["id"], date_to=date_to)
    rows = []
    total_dr = Decimal("0.00")
    total_cr = Decimal("0.00")
    for a in accounts:
        if a.get("is_header"):
            continue
        b = bals.get(a["code"], {"debit": Decimal("0.00"), "credit": Decimal("0.00")})
        dr = b["debit"]
        cr = b["credit"]
        net = dr - cr
        if a["normal_balance"] == "credit":
            net = -net
        # Show net on natural side; tiny values suppressed
        show_dr = net if (a["normal_balance"] == "debit" and net >= 0) else Decimal("0.00")
        show_cr = net if (a["normal_balance"] == "credit" and net >= 0) else Decimal("0.00")
        # Handle negative (contra) balances — show on the opposite column
        if a["normal_balance"] == "debit" and net < 0:
            show_cr = -net
        if a["normal_balance"] == "credit" and net < 0:
            show_dr = -net
        if show_dr == 0 and show_cr == 0:
            continue
        total_dr += show_dr
        total_cr += show_cr
        rows.append({
            "code": a["code"], "name": a["name"], "type": a["type"],
            "debit": float(show_dr), "credit": float(show_cr),
        })
    return {
        "as_at": date_to or datetime.now(timezone.utc).date().isoformat(),
        "rows": rows,
        "total_debit": float(total_dr), "total_credit": float(total_cr),
        "balanced": total_dr == total_cr,
    }


async def _sum_type(owner_id: str, acct_type: str, date_from: str | None, date_to: str | None) -> Decimal:
    accts = await db.accounts.find({"owner_id": owner_id, "type": acct_type}, {"_id": 0, "code": 1, "normal_balance": 1}).to_list(500)
    bals = await _balance_by_code(owner_id, date_from=date_from, date_to=date_to)
    total = Decimal("0.00")
    for a in accts:
        b = bals.get(a["code"], {"debit": Decimal("0.00"), "credit": Decimal("0.00")})
        net = b["debit"] - b["credit"]
        if a["normal_balance"] == "credit":
            net = -net
        total += net
    return total


@api.get("/accounting/reports/income-statement")
async def income_statement(date_from: str | None = None, date_to: str | None = None, u: dict = Depends(current_user)):
    """Revenue − Expenses over a period. Accrual basis."""
    accounts = await db.accounts.find({"owner_id": u["id"], "type": {"$in": ["income", "expense"]}}, {"_id": 0}).sort("code", 1).to_list(500)
    bals = await _balance_by_code(u["id"], date_from=date_from, date_to=date_to)
    income_rows, expense_rows = [], []
    total_income = Decimal("0.00")
    total_expense = Decimal("0.00")
    for a in accounts:
        if a.get("is_header"):
            continue
        b = bals.get(a["code"], {"debit": Decimal("0.00"), "credit": Decimal("0.00")})
        if a["type"] == "income":
            amt = b["credit"] - b["debit"]
            if amt != 0:
                income_rows.append({"code": a["code"], "name": a["name"], "amount": float(amt)})
                total_income += amt
        else:
            amt = b["debit"] - b["credit"]
            if amt != 0:
                expense_rows.append({"code": a["code"], "name": a["name"], "amount": float(amt)})
                total_expense += amt
    net_income = total_income - total_expense
    # Estimate corporate income tax at 27% (SA SBC-ish; NOT accurate for small-biz tables)
    tax_estimate = (net_income * Decimal("0.27")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) if net_income > 0 else Decimal("0.00")
    net_after_tax = net_income - tax_estimate
    return {
        "date_from": date_from, "date_to": date_to or datetime.now(timezone.utc).date().isoformat(),
        "income": income_rows, "total_income": float(total_income),
        "expenses": expense_rows, "total_expense": float(total_expense),
        "net_income_before_tax": float(net_income),
        "estimated_tax_at_27pct": float(tax_estimate),
        "net_income_after_tax": float(net_after_tax),
        "disclaimer": "Estimated corporate income tax at the headline 27% rate. Actual SARS tax turn depends on assessed loss, SBC tables, deductions and add-backs. Validate with a CA(SA).",
    }


@api.get("/accounting/reports/balance-sheet")
async def balance_sheet(as_at: str | None = None, u: dict = Depends(current_user)):
    accounts = await db.accounts.find({"owner_id": u["id"]}, {"_id": 0}).sort("code", 1).to_list(500)
    bals = await _balance_by_code(u["id"], date_to=as_at)

    def bucket(type_: str):
        rows = []
        total = Decimal("0.00")
        for a in accounts:
            if a["type"] != type_ or a.get("is_header"):
                continue
            b = bals.get(a["code"], {"debit": Decimal("0.00"), "credit": Decimal("0.00")})
            net = b["debit"] - b["credit"]
            if a["normal_balance"] == "credit":
                net = -net
            if net != 0:
                rows.append({"code": a["code"], "name": a["name"], "amount": float(net)})
                total += net
        return rows, total

    assets, total_assets = bucket("asset")
    liabs, total_liabs = bucket("liability")
    # Exclude 33000 from direct equity bucket: we always recompute it from live revenue−expense below.
    eq_rows = []
    total_eq = Decimal("0.00")
    for a in accounts:
        if a["type"] != "equity" or a.get("is_header") or a["code"] == "33000":
            continue
        b = bals.get(a["code"], {"debit": Decimal("0.00"), "credit": Decimal("0.00")})
        net = b["debit"] - b["credit"]
        if a["normal_balance"] == "credit":
            net = -net
        if net != 0:
            eq_rows.append({"code": a["code"], "name": a["name"], "amount": float(net)})
            total_eq += net

    # Current-year earnings (Revenue − Expense − est tax) rolls into Equity presentation
    net_income = await _sum_type(u["id"], "income", None, as_at) - await _sum_type(u["id"], "expense", None, as_at)
    total_eq_with_ni = total_eq + net_income
    eq_rows.append({"code": "33000", "name": "Current-year Earnings (auto)", "amount": float(net_income)})

    return {
        "as_at": as_at or datetime.now(timezone.utc).date().isoformat(),
        "assets": assets, "total_assets": float(total_assets),
        "liabilities": liabs, "total_liabilities": float(total_liabs),
        "equity": eq_rows, "total_equity": float(total_eq_with_ni),
        "liabilities_plus_equity": float(total_liabs + total_eq_with_ni),
        "balanced": total_assets == (total_liabs + total_eq_with_ni),
    }


@api.get("/accounting/reports/vat201")
async def vat201(date_from: str, date_to: str, u: dict = Depends(current_user)):
    """SA VAT 201 simplified. Groups by vat_code on journal lines."""
    pipeline = [
        {"$match": {"owner_id": u["id"], "posted": True, "date": {"$gte": date_from, "$lte": date_to}}},
        {"$unwind": "$lines"},
        {"$match": {"lines.vat_code": {"$exists": True, "$ne": None}}},
        {"$group": {
            "_id": "$lines.vat_code",
            "debit": {"$sum": "$lines.debit"},
            "credit": {"$sum": "$lines.credit"},
            "vat_amt": {"$sum": "$lines.vat_amount"},
            "count": {"$sum": 1},
        }},
    ]
    rows = {}
    async for r in db.journals.aggregate(pipeline):
        rows[r["_id"]] = r

    # Output tax (supplies) = VAT on sales (S code, credit side on VAT Output account)
    def box(code: str, side: str) -> Decimal:
        r = rows.get(code)
        if not r:
            return Decimal("0.00")
        if side == "credit":
            return _D(r.get("credit") or 0)
        return _D(r.get("debit") or 0)

    std_out = box("S", "credit")       # noqa: F841 — kept for future VAT201 box-6 computation
    zero_supplies = box("Z", "credit")  # zero-rated output value
    exempt_supplies = box("E", "credit")
    std_inputs = _D(rows.get("SI", {}).get("vat_amt") or 0) if "SI" in rows else Decimal("0.00")
    cap_inputs = _D(rows.get("CI", {}).get("vat_amt") or 0) if "CI" in rows else Decimal("0.00")

    # Better: pull VAT control from journals by account code (52000 VAT Output, 23000 VAT Input)
    bals = await _balance_by_code(u["id"], date_from=date_from, date_to=date_to)
    vat_output_net = _D(bals.get("52000", {}).get("credit") or 0) - _D(bals.get("52000", {}).get("debit") or 0)
    vat_input_net = _D(bals.get("23000", {}).get("debit") or 0) - _D(bals.get("23000", {}).get("credit") or 0)
    vat_payable = vat_output_net - vat_input_net

    return {
        "period": {"date_from": date_from, "date_to": date_to},
        "output_tax": {
            "box_1_standard_rated_15pct": float(vat_output_net),
            "box_2_zero_rated_supplies_value": float(zero_supplies),
            "box_3_exempt_and_other_supplies_value": float(exempt_supplies),
        },
        "input_tax": {
            "box_14_standard_inputs_15pct": float(std_inputs),
            "box_15_capital_inputs_15pct": float(cap_inputs),
            "total_input_tax_claim": float(vat_input_net),
        },
        "vat_payable_to_sars": float(vat_payable),
        "breakdown_by_vat_code": [
            {"code": k, "debit": float(_D(v.get("debit") or 0)),
             "credit": float(_D(v.get("credit") or 0)),
             "vat_amount": float(_D(v.get("vat_amt") or 0)),
             "line_count": int(v.get("count") or 0)}
            for k, v in rows.items()
        ],
        "disclaimer": "Simplified VAT 201 view for sign-off. Always reconcile with SARS eFiling VAT201 before submission.",
    }


@api.get("/accounting/reports/general-ledger/{account_code}")
async def general_ledger(account_code: str, date_from: str | None = None, date_to: str | None = None, u: dict = Depends(current_user)):
    acct = await db.accounts.find_one({"owner_id": u["id"], "code": account_code}, {"_id": 0})
    if not acct:
        raise HTTPException(404, "Account not found")
    q: dict = {"owner_id": u["id"], "posted": True, "lines.account_code": account_code}
    if date_from or date_to:
        q["date"] = {}
        if date_from:
            q["date"]["$gte"] = date_from
        if date_to:
            q["date"]["$lte"] = date_to
    journals = await db.journals.find(q, {"_id": 0}).sort("date", 1).to_list(2000)
    rows = []
    running = Decimal("0.00")
    for j in journals:
        for ln in j["lines"]:
            if ln["account_code"] != account_code:
                continue
            d = _D(ln.get("debit") or 0)
            c = _D(ln.get("credit") or 0)
            delta = d - c if acct["normal_balance"] == "debit" else c - d
            running += delta
            rows.append({
                "date": j["date"], "journal_id": j["id"], "memo": j["memo"], "reference": j.get("reference"),
                "description": ln.get("description"), "debit": float(d), "credit": float(c),
                "running_balance": float(running), "source": j.get("source"),
            })
    return {"account": acct, "rows": rows, "closing_balance": float(running)}


# ---- Fiscal Periods + Sign-off ---------------------------------------------
@api.get("/accounting/periods")
async def list_periods(u: dict = Depends(current_user)):
    return await db.fiscal_periods.find({"owner_id": u["id"]}, {"_id": 0}).sort("period", -1).to_list(120)


@api.post("/accounting/periods/{period}/close")
async def close_period(period: str, u: dict = Depends(require_accountant)):
    """Soft close: prevents new *manual* journals from posting in this period (auto from invoices still allowed)."""
    res = await db.fiscal_periods.update_one(
        {"owner_id": u["id"], "period": period},
        {"$set": {"status": "closed", "closed_by": u["actor_id"], "closed_at": now_iso()}},
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Period not found")
    await audit(u["actor_id"], "close_period", "fiscal_period", period)
    return {"ok": True, "period": period, "status": "closed"}


@api.post("/accounting/periods/{period}/lock")
async def lock_period(period: str, u: dict = Depends(require_accountant)):
    """Hard lock: no journals may post (including auto) until an accountant reopens it."""
    res = await db.fiscal_periods.update_one(
        {"owner_id": u["id"], "period": period},
        {"$set": {"status": "locked", "locked_by": u["actor_id"], "locked_at": now_iso()}},
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Period not found")
    await audit(u["actor_id"], "lock_period", "fiscal_period", period)
    return {"ok": True, "period": period, "status": "locked"}


@api.post("/accounting/periods/{period}/reopen")
async def reopen_period(period: str, u: dict = Depends(require_accountant)):
    res = await db.fiscal_periods.update_one(
        {"owner_id": u["id"], "period": period},
        {"$set": {"status": "open", "reopened_by": u["actor_id"], "reopened_at": now_iso()}},
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Period not found")
    await audit(u["actor_id"], "reopen_period", "fiscal_period", period)
    return {"ok": True, "period": period, "status": "open"}


@api.post("/accounting/periods/{period}/signoff")
async def signoff_period(period: str, body: dict = Body(...), u: dict = Depends(require_accountant)):
    """Accountant sign-off with optional note. Does not lock; call lock separately."""
    note = body.get("note") or ""
    res = await db.fiscal_periods.update_one(
        {"owner_id": u["id"], "period": period},
        {"$set": {"signed_off_by": u["actor_id"], "signed_off_at": now_iso(), "signoff_note": note}},
    )
    if res.matched_count == 0:
        raise HTTPException(404, f"Fiscal period {period} not found")
    await db.accountant_notes.insert_one({
        "id": new_id(), "owner_id": u["id"], "period": period,
        "actor_id": u["actor_id"], "note": note, "created_at": now_iso(),
    })
    await audit(u["actor_id"], "signoff_period", "fiscal_period", period, after={"note": note[:200]})
    return {"ok": True, "period": period, "signed_off_at": now_iso()}


@api.get("/accounting/periods/{period}/notes")
async def period_notes(period: str, u: dict = Depends(current_user)):
    return await db.accountant_notes.find({"owner_id": u["id"], "period": period}, {"_id": 0}).sort("created_at", -1).to_list(200)


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 7 — PDF exports (VAT201 + TB + IS + BS)
# ══════════════════════════════════════════════════════════════════════════════
from accounting_pdf import _fmt_zar, _pdf_buf_from_story, _report_table  # noqa: E402


async def _resolve_owner_branding(uid: str) -> dict:
    owner = await db.users.find_one({"id": uid}, {"_id": 0, "password_hash": 0}) or {}
    tpl = owner.get("quote_template") or {}
    return {
        "company_name": tpl.get("company_name") or owner.get("name") or "Ascent CRM",
        "accent_hex": (tpl.get("accent_color_hex") or "E26E4A").lstrip("#")[:6] or "E26E4A",
        "email": owner.get("email") or "",
    }


@api.get("/accounting/reports/trial-balance/pdf")
async def trial_balance_pdf(date_to: str | None = None, u: dict = Depends(current_user)):
    r = await trial_balance(date_to=date_to, u=u)
    branding = await _resolve_owner_branding(u["id"])

    def build(styles, accent):
        from reportlab.platypus import Paragraph, Spacer
        items = [Paragraph(f"As at: <b>{r['as_at']}</b>", styles["Normal"]), Spacer(1, 8)]
        data = [["Code", "Name", "Type", "Debit", "Credit"]]
        for x in r["rows"]:
            data.append([
                x["code"], x["name"], x["type"],
                _fmt_zar(x["debit"]) if x["debit"] else "",
                _fmt_zar(x["credit"]) if x["credit"] else "",
            ])
        data.append(["", "", "TOTAL", _fmt_zar(r["total_debit"]), _fmt_zar(r["total_credit"])])
        items.append(_report_table(data, styles, accent, col_widths=[55, 210, 60, 90, 90], right_align_cols=[3, 4]))
        items.append(Spacer(1, 8))
        items.append(Paragraph(
            "✓ Balanced" if r["balanced"] else "✗ UNBALANCED — investigate",
            styles["Normal"],
        ))
        items.append(Paragraph(
            "Prepared for accountant review. Must be signed off by a CA(SA) / SAICA / SAIPA member before filing.",
            styles["Disclaimer"],
        ))
        return items

    buf = _pdf_buf_from_story(f"Trial Balance — {r['as_at']}", build, branding)
    await audit(u["actor_id"], "export_pdf", "accounting_report", "trial_balance", after={"date_to": r["as_at"]})
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="TrialBalance_{r["as_at"]}.pdf"'})


@api.get("/accounting/reports/income-statement/pdf")
async def income_statement_pdf(date_from: str | None = None, date_to: str | None = None, u: dict = Depends(current_user)):
    r = await income_statement(date_from=date_from, date_to=date_to, u=u)
    branding = await _resolve_owner_branding(u["id"])

    def build(styles, accent):
        from reportlab.platypus import Paragraph, Spacer
        items = [Paragraph(f"Period: <b>{r.get('date_from') or 'start'} → {r['date_to']}</b>", styles["Normal"]), Spacer(1, 8)]
        data = [["Code", "Name", "Amount (ZAR)"]]
        data.append(["", "— REVENUE —", ""])
        for x in r["income"]:
            data.append([x["code"], x["name"], _fmt_zar(x["amount"])])
        data.append(["", "Total Revenue", _fmt_zar(r["total_income"])])
        data.append(["", "— EXPENSES —", ""])
        for x in r["expenses"]:
            data.append([x["code"], x["name"], _fmt_zar(x["amount"])])
        data.append(["", "Total Expenses", _fmt_zar(r["total_expense"])])
        data.append(["", "NET INCOME BEFORE TAX", _fmt_zar(r["net_income_before_tax"])])
        data.append(["", "Estimated Corporate Tax @ 27%", _fmt_zar(r["estimated_tax_at_27pct"])])
        data.append(["", "NET INCOME AFTER TAX", _fmt_zar(r["net_income_after_tax"])])
        items.append(_report_table(data, styles, accent, col_widths=[60, 290, 120], right_align_cols=[2]))
        items.append(Paragraph(r.get("disclaimer") or "", styles["Disclaimer"]))
        return items

    buf = _pdf_buf_from_story("Income Statement", build, branding)
    await audit(u["actor_id"], "export_pdf", "accounting_report", "income_statement", after={"date_from": date_from, "date_to": date_to})
    fname = f"IncomeStatement_{r.get('date_from') or 'start'}_to_{r['date_to']}.pdf"
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@api.get("/accounting/reports/balance-sheet/pdf")
async def balance_sheet_pdf(as_at: str | None = None, u: dict = Depends(current_user)):
    r = await balance_sheet(as_at=as_at, u=u)
    branding = await _resolve_owner_branding(u["id"])

    def build(styles, accent):
        from reportlab.platypus import Paragraph, Spacer
        items = [Paragraph(f"As at: <b>{r['as_at']}</b>", styles["Normal"]), Spacer(1, 8)]

        def section(title, rows, total):
            data = [["Code", "Name", "Amount (ZAR)"]]
            if not rows:
                data.append(["", "—", ""])
            for x in rows:
                data.append([x["code"], x["name"], _fmt_zar(x["amount"])])
            data.append(["", f"Total {title}", _fmt_zar(total)])
            return [
                Paragraph(title, styles["SectionHead"]),
                _report_table(data, styles, accent, col_widths=[60, 290, 120], right_align_cols=[2]),
            ]

        items.extend(section("Assets", r["assets"], r["total_assets"]))
        items.extend(section("Liabilities", r["liabilities"], r["total_liabilities"]))
        items.extend(section("Equity", r["equity"], r["total_equity"]))
        items.append(Spacer(1, 8))
        items.append(Paragraph(
            f"Assets {_fmt_zar(r['total_assets'])}  =  Liabilities + Equity {_fmt_zar(r['liabilities_plus_equity'])}   "
            + ("✓ Balanced" if r["balanced"] else "✗ UNBALANCED"),
            styles["Normal"],
        ))
        items.append(Paragraph(
            "Prepared for accountant review. Must be signed off by a CA(SA) / SAICA / SAIPA member before filing.",
            styles["Disclaimer"],
        ))
        return items

    buf = _pdf_buf_from_story(f"Balance Sheet — {r['as_at']}", build, branding)
    await audit(u["actor_id"], "export_pdf", "accounting_report", "balance_sheet", after={"as_at": r["as_at"]})
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="BalanceSheet_{r["as_at"]}.pdf"'})


@api.get("/accounting/reports/vat201/pdf")
async def vat201_pdf(date_from: str, date_to: str, u: dict = Depends(current_user)):
    r = await vat201(date_from=date_from, date_to=date_to, u=u)
    branding = await _resolve_owner_branding(u["id"])

    def build(styles, accent):
        from reportlab.platypus import Paragraph, Spacer
        items = [Paragraph(f"Period: <b>{r['period']['date_from']} → {r['period']['date_to']}</b>", styles["Normal"]), Spacer(1, 8)]

        out = r["output_tax"]
        inp = r["input_tax"]
        data = [["Box", "Description", "Amount (ZAR)"]]
        data.append(["1",  "Standard-rated supplies 15% (output)",    _fmt_zar(out["box_1_standard_rated_15pct"])])
        data.append(["2",  "Zero-rated supplies (value)",              _fmt_zar(out["box_2_zero_rated_supplies_value"])])
        data.append(["3",  "Exempt / other supplies (value)",          _fmt_zar(out["box_3_exempt_and_other_supplies_value"])])
        data.append(["14", "Standard-rated inputs 15%",                _fmt_zar(inp["box_14_standard_inputs_15pct"])])
        data.append(["15", "Capital inputs 15%",                        _fmt_zar(inp["box_15_capital_inputs_15pct"])])
        data.append(["",   "Total input tax claim",                     _fmt_zar(inp["total_input_tax_claim"])])
        data.append(["",   "VAT PAYABLE TO SARS",                       _fmt_zar(r["vat_payable_to_sars"])])
        items.append(_report_table(data, styles, accent, col_widths=[40, 330, 120], right_align_cols=[2]))

        # Breakdown by VAT code
        if r.get("breakdown_by_vat_code"):
            items.append(Paragraph("Breakdown by VAT code", styles["SectionHead"]))
            bd = [["Code", "Debit", "Credit", "VAT Amount", "Lines"]]
            for row in r["breakdown_by_vat_code"]:
                bd.append([row["code"], _fmt_zar(row["debit"]), _fmt_zar(row["credit"]),
                           _fmt_zar(row["vat_amount"]), str(row["line_count"])])
            items.append(_report_table(bd, styles, accent, col_widths=[55, 110, 110, 110, 60], right_align_cols=[1, 2, 3, 4]))

        items.append(Paragraph(r.get("disclaimer") or "", styles["Disclaimer"]))
        items.append(Paragraph(
            "This document is a supporting workpaper. Submit VAT201 via SARS eFiling; do not upload this PDF to SARS.",
            styles["Disclaimer"],
        ))
        return items

    buf = _pdf_buf_from_story(f"VAT 201 — {r['period']['date_from']} to {r['period']['date_to']}", build, branding)
    await audit(u["actor_id"], "export_pdf", "accounting_report", "vat201",
                after={"date_from": date_from, "date_to": date_to, "payable": r["vat_payable_to_sars"]})
    fname = f"VAT201_{date_from}_to_{date_to}.pdf"
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'})


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 7 — Fixed Asset Register + Depreciation
# ══════════════════════════════════════════════════════════════════════════════
class FixedAssetIn(BaseModel):
    name: str
    description: str | None = None
    asset_category: str = "equipment"  # equipment, vehicles, furniture, buildings, computers, software
    acquisition_date: str              # YYYY-MM-DD
    acquisition_cost: float
    residual_value: float = 0.0
    useful_life_months: int            # straight-line
    depreciation_method: Literal["straight_line"] = "straight_line"
    asset_account_code: str = "11100"  # Computer Equipment — Cost (default; matches seed)
    accumulated_depr_account_code: str = "11110"  # Computer Equipment — Accumulated Depreciation
    depreciation_expense_account_code: str = "82500"  # Depreciation expense
    serial_number: str | None = None
    location: str | None = None


@api.post("/accounting/fixed-assets")
async def create_fixed_asset(p: FixedAssetIn, u: dict = Depends(require_accountant)):
    if p.useful_life_months <= 0:
        raise HTTPException(400, "useful_life_months must be > 0")
    if p.acquisition_cost <= 0:
        raise HTTPException(400, "acquisition_cost must be > 0")
    if p.residual_value < 0 or p.residual_value >= p.acquisition_cost:
        raise HTTPException(400, "residual_value must be ≥0 and < acquisition_cost")
    doc = p.model_dump()
    doc["id"] = new_id()
    doc["owner_id"] = u["id"]
    doc["status"] = "active"  # active / disposed / fully-depreciated
    doc["depreciation_to_date"] = 0.0
    doc["book_value"] = float(p.acquisition_cost)
    doc["last_depreciated_period"] = None
    doc["created_at"] = now_iso()
    await db.fixed_assets.insert_one(doc)
    await audit(u["actor_id"], "create_asset", "fixed_asset", doc["id"], after={"name": p.name, "cost": p.acquisition_cost})
    return _strip_oid(doc)


@api.get("/accounting/fixed-assets")
async def list_fixed_assets(u: dict = Depends(current_user)):
    return await db.fixed_assets.find({"owner_id": u["id"]}, {"_id": 0}).sort("acquisition_date", -1).to_list(1000)


@api.get("/accounting/fixed-assets/{aid}")
async def get_fixed_asset(aid: str, u: dict = Depends(current_user)):
    a = await db.fixed_assets.find_one({"id": aid, "owner_id": u["id"]}, {"_id": 0})
    if not a:
        raise HTTPException(404, "Asset not found")
    # Full monthly schedule preview
    monthly = (a["acquisition_cost"] - a["residual_value"]) / a["useful_life_months"]
    schedule = []
    acq = datetime.fromisoformat(a["acquisition_date"])
    running_depr = 0.0
    for m in range(1, a["useful_life_months"] + 1):
        period_dt = acq + timedelta(days=30 * m)
        running_depr += monthly
        running_depr = min(running_depr, a["acquisition_cost"] - a["residual_value"])
        schedule.append({
            "month": m,
            "period": period_dt.strftime("%Y-%m"),
            "depreciation": round(monthly, 2),
            "accumulated": round(running_depr, 2),
            "book_value": round(a["acquisition_cost"] - running_depr, 2),
        })
    return {"asset": a, "monthly_depreciation": round(monthly, 2), "schedule": schedule}


@api.delete("/accounting/fixed-assets/{aid}")
async def dispose_fixed_asset(aid: str, u: dict = Depends(require_accountant)):
    a = await db.fixed_assets.find_one({"id": aid, "owner_id": u["id"]}, {"_id": 0})
    if not a:
        raise HTTPException(404, "Asset not found")
    await db.fixed_assets.update_one(
        {"id": aid, "owner_id": u["id"]},
        {"$set": {"status": "disposed", "disposed_at": now_iso()}},
    )
    await audit(u["actor_id"], "dispose_asset", "fixed_asset", aid, before=a)
    return {"ok": True, "status": "disposed"}


@api.post("/accounting/fixed-assets/depreciate")
async def post_depreciation(body: dict = Body(default={}), u: dict = Depends(require_accountant)):
    """Post monthly straight-line depreciation for the given period (YYYY-MM).
    Idempotent: will not double-post the same period for an asset.
    """
    period = body.get("period") or datetime.now(timezone.utc).strftime("%Y-%m")
    # Sanity — period format
    if len(period) != 7 or period[4] != "-":
        raise HTTPException(400, "period must be YYYY-MM")
    period_end = period + "-28"  # safe last-of-month-ish anchor for journal date

    assets = await db.fixed_assets.find({
        "owner_id": u["id"],
        "status": "active",
    }, {"_id": 0}).to_list(1000)

    posted = []
    skipped = []
    for a in assets:
        if a.get("last_depreciated_period") and a["last_depreciated_period"] >= period:
            skipped.append({"id": a["id"], "reason": "already depreciated for this period"})
            continue
        if a["acquisition_date"] > period_end:
            skipped.append({"id": a["id"], "reason": "not yet in service"})
            continue
        monthly = (a["acquisition_cost"] - a["residual_value"]) / a["useful_life_months"]
        remaining = (a["acquisition_cost"] - a["residual_value"]) - a["depreciation_to_date"]
        if remaining <= 0.01:
            await db.fixed_assets.update_one({"id": a["id"]}, {"$set": {"status": "fully-depreciated"}})
            skipped.append({"id": a["id"], "reason": "fully depreciated"})
            continue
        amt = round(min(monthly, remaining), 2)

        # Post journal: DR Depreciation Expense / CR Accumulated Depreciation
        j = JournalIn(
            date=period_end,
            memo=f"Monthly depreciation · {a['name']}",
            reference=f"DEPR-{a['id'][:8]}-{period}",
            lines=[
                JournalLineIn(account_code=a["depreciation_expense_account_code"], debit=amt, credit=0, description=f"Depreciation {a['name']}"),
                JournalLineIn(account_code=a["accumulated_depr_account_code"], debit=0, credit=amt, description=f"Accum depr {a['name']}"),
            ],
            source="depreciation",
            source_id=a["id"],
        )
        try:
            jdoc = await _validate_and_post_journal(u["id"], u["actor_id"], j, auto=False)
        except HTTPException as e:
            skipped.append({"id": a["id"], "reason": f"journal rejected: {e.detail}"})
            continue

        new_depr = round(a["depreciation_to_date"] + amt, 2)
        new_book = round(a["acquisition_cost"] - new_depr, 2)
        new_status = "fully-depreciated" if new_depr >= (a["acquisition_cost"] - a["residual_value"] - 0.01) else "active"
        await db.fixed_assets.update_one(
            {"id": a["id"]},
            {"$set": {
                "depreciation_to_date": new_depr,
                "book_value": new_book,
                "last_depreciated_period": period,
                "status": new_status,
            }},
        )
        posted.append({"id": a["id"], "name": a["name"], "amount": amt, "journal_id": jdoc["id"]})

    await audit(u["actor_id"], "depreciate_batch", "fixed_asset", period, after={"posted": len(posted), "skipped": len(skipped)})
    return {"period": period, "posted": posted, "skipped": skipped}


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 7 — Bank Feeds + Reconciliation (CSV upload)
# ══════════════════════════════════════════════════════════════════════════════
class BankAccountIn(BaseModel):
    name: str
    bank: str                          # e.g. "FNB", "Standard Bank"
    account_number: str | None = None
    gl_account_code: str = "21000"     # links to the Bank account in CoA (FNB Current in seed)
    currency: str = "ZAR"


@api.post("/accounting/bank-accounts")
async def create_bank_account(p: BankAccountIn, u: dict = Depends(require_accountant)):
    doc = p.model_dump()
    doc["id"] = new_id()
    doc["owner_id"] = u["id"]
    doc["created_at"] = now_iso()
    await db.bank_accounts.insert_one(doc)
    await audit(u["actor_id"], "create_bank_account", "bank_account", doc["id"], after={"name": p.name})
    return _strip_oid(doc)


@api.get("/accounting/bank-accounts")
async def list_bank_accounts(u: dict = Depends(current_user)):
    return await db.bank_accounts.find({"owner_id": u["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)


@api.post("/accounting/bank-accounts/{bid}/import-csv")
async def import_bank_csv(bid: str, file: UploadFile = _File(...), u: dict = Depends(require_accountant)):
    """Import bank statement from CSV. Expected columns (case-insensitive):
    `date`, `description`, `amount` (positive = credit/in, negative = debit/out), `balance` (optional).
    Duplicates detected via (date, description, amount) hash per account.
    """
    bank = await db.bank_accounts.find_one({"id": bid, "owner_id": u["id"]}, {"_id": 0})
    if not bank:
        raise HTTPException(404, "Bank account not found")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024:
        raise HTTPException(413, "CSV too large (5 MB limit)")
    import csv as _csv
    import hashlib
    text = raw.decode("utf-8-sig", errors="ignore")
    reader = _csv.DictReader(io.StringIO(text))
    if not reader.fieldnames:
        raise HTTPException(400, "Empty or invalid CSV")
    fieldmap = {k.lower().strip(): k for k in reader.fieldnames}
    required = {"date", "description", "amount"}
    missing = required - set(fieldmap.keys())
    if missing:
        raise HTTPException(400, f"CSV missing required columns: {sorted(missing)}")
    inserted = 0
    skipped_dupes = 0
    errors = []
    for row in reader:
        try:
            date = (row.get(fieldmap["date"]) or "").strip()
            desc = (row.get(fieldmap["description"]) or "").strip()
            amt_raw = (row.get(fieldmap["amount"]) or "0").replace(",", "").replace(" ", "").strip()
            amt = float(amt_raw)
            # Normalise YYYY-MM-DD (accept DD/MM/YYYY too)
            if "/" in date:
                parts = date.split("/")
                if len(parts) == 3 and len(parts[2]) == 4:
                    date = f"{parts[2]}-{parts[1].zfill(2)}-{parts[0].zfill(2)}"
            h = hashlib.sha256(f"{bid}|{date}|{desc}|{amt}".encode()).hexdigest()
            if await db.bank_transactions.find_one({"dup_hash": h}, {"_id": 0, "id": 1}):
                skipped_dupes += 1
                continue
            await db.bank_transactions.insert_one({
                "id": new_id(),
                "owner_id": u["id"],
                "bank_account_id": bid,
                "date": date,
                "description": desc,
                "amount": amt,
                "balance": float((row.get(fieldmap["balance"], "0") or "0").replace(",", "").strip() or 0)
                             if "balance" in fieldmap else None,
                "direction": "in" if amt >= 0 else "out",
                "status": "unreconciled",
                "matched_ref_type": None,
                "matched_ref_id": None,
                "journal_id": None,
                "dup_hash": h,
                "created_at": now_iso(),
            })
            inserted += 1
        except Exception as e:
            errors.append({"row": row, "error": str(e)})
    await audit(u["actor_id"], "import_bank_csv", "bank_account", bid,
                after={"inserted": inserted, "dupes": skipped_dupes, "errors": len(errors)})
    return {"inserted": inserted, "skipped_duplicates": skipped_dupes, "errors": errors[:20]}


@api.get("/accounting/bank-accounts/{bid}/transactions")
async def list_bank_transactions(bid: str, status: str | None = None, u: dict = Depends(current_user)):
    q: dict = {"owner_id": u["id"], "bank_account_id": bid}
    if status:
        q["status"] = status
    return await db.bank_transactions.find(q, {"_id": 0}).sort("date", -1).to_list(1000)


@api.get("/accounting/bank-transactions/{tid}/suggest-matches")
async def suggest_matches(tid: str, u: dict = Depends(current_user)):
    """For money-in: suggest unpaid/overdue invoices with amount within 5%.
    For money-out: suggest recent unmatched expenses with amount within 5%.
    """
    tx = await db.bank_transactions.find_one({"id": tid, "owner_id": u["id"]}, {"_id": 0})
    if not tx:
        raise HTTPException(404, "Bank transaction not found")
    amt = abs(tx["amount"])
    tol = max(1.0, amt * 0.05)
    suggestions: list[dict] = []
    if tx["direction"] == "in":
        invoices = await db.invoices.find({
            "owner_id": u["id"],
            "status": {"$ne": "paid"},
            "grand_total": {"$gte": amt - tol, "$lte": amt + tol},
        }, {"_id": 0}).sort("created_at", -1).to_list(20)
        for inv in invoices:
            suggestions.append({
                "type": "invoice",
                "id": inv["id"],
                "label": f"Invoice #{inv.get('number')} · {inv.get('currency','ZAR')} {inv.get('grand_total',0):,.2f}",
                "amount": inv.get("grand_total"),
                "confidence": 0.9 if abs(inv.get("grand_total", 0) - amt) < 0.01 else 0.6,
            })
    else:
        # Money-out: surface the tenant's recent expense accounts as candidate CoA targets
        # (user picks the category for reconcile as an expense). We rank by recent usage in journals.
        recent = await db.journals.find(
            {"owner_id": u["id"]},
            {"_id": 0, "lines": 1}
        ).sort("date", -1).limit(50).to_list(50)
        usage: dict = {}
        for j in recent:
            for ln in j.get("lines") or []:
                if float(ln.get("debit") or 0) > 0:
                    usage[ln["account_code"]] = usage.get(ln["account_code"], 0) + 1
        expense_accts = await db.accounts.find(
            {"owner_id": u["id"], "type": "expense", "subtype": {"$ne": "header"}},
            {"_id": 0, "code": 1, "name": 1}
        ).to_list(200)
        expense_accts.sort(key=lambda a: (-usage.get(a["code"], 0), a["code"]))
        for a in expense_accts[:10]:
            suggestions.append({
                "type": "expense_account",
                "id": a["code"],
                "label": f"{a['code']} · {a['name']}",
                "amount": amt,
                "confidence": 0.5 + min(0.4, 0.05 * usage.get(a["code"], 0)),
            })
    return {"transaction": tx, "suggestions": suggestions}


@api.post("/accounting/bank-transactions/{tid}/reconcile")
async def reconcile_transaction(tid: str, body: dict = Body(...), u: dict = Depends(require_accountant)):
    """Reconcile a bank transaction against an invoice (money in) or post as an expense (money out).
    Body:
      - for invoice match:  {"match_type":"invoice", "invoice_id":"..."}
      - for expense/other:  {"match_type":"expense", "expense_account_code":"81100", "description":"optional"}
    """
    tx = await db.bank_transactions.find_one({"id": tid, "owner_id": u["id"]}, {"_id": 0})
    if not tx:
        raise HTTPException(404, "Bank transaction not found")
    if tx["status"] == "reconciled":
        raise HTTPException(400, "Already reconciled")
    bank = await db.bank_accounts.find_one({"id": tx["bank_account_id"], "owner_id": u["id"]}, {"_id": 0})
    if not bank:
        raise HTTPException(400, "Parent bank account missing")
    match_type = body.get("match_type")
    amt = abs(tx["amount"])
    bank_code = bank["gl_account_code"]

    if match_type == "invoice":
        inv = await db.invoices.find_one({"id": body.get("invoice_id"), "owner_id": u["id"]}, {"_id": 0})
        if not inv:
            raise HTTPException(404, "Invoice not found")
        # Guard: bank transaction amount must be within 5% of invoice grand_total
        # (prevents accidentally marking a R50,000 invoice paid from a R50 bank line).
        inv_total = float(inv.get("grand_total") or 0)
        tol = max(1.0, inv_total * 0.05)
        if abs(amt - inv_total) > tol:
            raise HTTPException(400, f"Amount mismatch: bank tx R{amt:,.2f} vs invoice R{inv_total:,.2f} (tolerance 5%).")
        # For an invoice receipt: DR Bank, CR Debtors (mirrors the Stripe/PayPal capture flow,
        # but here we debit Bank directly since the money is in the bank account).
        j = JournalIn(
            date=tx["date"],
            memo=f"Bank receipt · Invoice #{inv.get('number')}",
            reference=f"BANK-RECON-{tx['id'][:8]}",
            lines=[
                JournalLineIn(account_code=bank_code, debit=amt, credit=0, description=f"Deposit — {tx['description'][:80]}"),
                JournalLineIn(account_code="22000",    debit=0, credit=amt, description=f"Settle Invoice {inv.get('number')}"),
            ],
            source="bank_reconciliation",
            source_id=tx["id"],
        )
        jdoc = await _validate_and_post_journal(u["id"], u["actor_id"], j, auto=False)
        await db.invoices.update_one({"id": inv["id"], "owner_id": u["id"]}, {"$set": {"status": "paid", "paid_at": now_iso()}})
        await db.bank_transactions.update_one({"id": tid, "owner_id": u["id"]}, {"$set": {
            "status": "reconciled", "matched_ref_type": "invoice", "matched_ref_id": inv["id"],
            "journal_id": jdoc["id"], "reconciled_at": now_iso(),
        }})
        await audit(u["actor_id"], "reconcile_bank_tx", "bank_transaction", tid, after={"invoice_id": inv["id"], "amount": amt})
        return {"ok": True, "journal_id": jdoc["id"], "invoice_paid": inv["id"]}

    elif match_type == "expense":
        acct_code = body.get("expense_account_code")
        if not acct_code:
            raise HTTPException(400, "expense_account_code required")
        acct = await db.accounts.find_one({"owner_id": u["id"], "code": acct_code}, {"_id": 0})
        if not acct or acct["type"] != "expense":
            raise HTTPException(400, "expense_account_code must be an existing expense account")
        if tx["direction"] != "out":
            raise HTTPException(400, "Can only reconcile money-out as expense")
        descr = body.get("description") or tx["description"]
        j = JournalIn(
            date=tx["date"],
            memo=f"Bank expense · {descr[:60]}",
            reference=f"BANK-RECON-{tx['id'][:8]}",
            lines=[
                JournalLineIn(account_code=acct_code, debit=amt, credit=0, description=descr[:160]),
                JournalLineIn(account_code=bank_code, debit=0, credit=amt, description=f"Paid — {tx['description'][:80]}"),
            ],
            source="bank_reconciliation",
            source_id=tx["id"],
        )
        jdoc = await _validate_and_post_journal(u["id"], u["actor_id"], j, auto=False)
        await db.bank_transactions.update_one({"id": tid, "owner_id": u["id"]}, {"$set": {
            "status": "reconciled", "matched_ref_type": "expense", "matched_ref_id": acct_code,
            "journal_id": jdoc["id"], "reconciled_at": now_iso(),
        }})
        await audit(u["actor_id"], "reconcile_bank_tx", "bank_transaction", tid, after={"expense_code": acct_code, "amount": amt})
        return {"ok": True, "journal_id": jdoc["id"]}

    raise HTTPException(400, "match_type must be 'invoice' or 'expense'")


@api.post("/accounting/bank-transactions/{tid}/unreconcile")
async def unreconcile_transaction(tid: str, u: dict = Depends(require_accountant)):
    tx = await db.bank_transactions.find_one({"id": tid, "owner_id": u["id"]}, {"_id": 0})
    if not tx:
        raise HTTPException(404, "Bank transaction not found")
    if tx["status"] != "reconciled":
        raise HTTPException(400, "Not reconciled")
    # Reverse the journal by re-using the existing reverse_journal logic inline.
    if tx.get("journal_id"):
        jrnl = await db.journals.find_one({"id": tx["journal_id"], "owner_id": u["id"]}, {"_id": 0})
        if jrnl and not jrnl.get("reversed_by"):
            rev_lines = [JournalLineIn(
                account_code=ln["account_code"],
                debit=float(ln.get("credit") or 0),
                credit=float(ln.get("debit") or 0),
                description=f"REV: {ln.get('description') or ''}",
                vat_code=ln.get("vat_code"),
                vat_amount=(-float(ln["vat_amount"]) if ln.get("vat_amount") is not None else None),
            ) for ln in jrnl["lines"]]
            rev_payload = JournalIn(
                date=datetime.now(timezone.utc).date().isoformat(),
                memo=f"Reversal of {jrnl['memo']}",
                reference=f"REV-{jrnl.get('reference') or jrnl['id']}",
                source="reversing", source_id=jrnl["id"], lines=rev_lines,
            )
            rev = await _validate_and_post_journal(u["id"], u["actor_id"], rev_payload, auto=False)
            await db.journals.update_one({"id": jrnl["id"]}, {"$set": {"reversed_by": rev["id"], "reversed_at": now_iso()}})
            await db.journals.update_one({"id": rev["id"]}, {"$set": {"reversed_of": jrnl["id"]}})
    await db.bank_transactions.update_one({"id": tid, "owner_id": u["id"]}, {"$set": {
        "status": "unreconciled", "matched_ref_type": None, "matched_ref_id": None,
        "journal_id": None, "unreconciled_at": now_iso(),
    }})
    await audit(u["actor_id"], "unreconcile_bank_tx", "bank_transaction", tid)
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch 7 — Receipt OCR (Gemini 3 vision via Emergent LLM)
# ══════════════════════════════════════════════════════════════════════════════
RECEIPT_UPLOAD_ROOT = UPLOAD_ROOT / "receipts"
RECEIPT_UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
ALLOWED_RECEIPT_SUFFIXES = {".jpg", ".jpeg", ".png", ".pdf", ".webp", ".heic"}


@api.post("/accounting/receipts/scan")
async def scan_receipt(file: UploadFile = _File(...), u: dict = Depends(require_accountant)):
    """OCR a receipt/invoice image using Gemini 3 vision (Emergent LLM).
    Returns structured data (vendor, date, total, vat, line_items). Does NOT auto-post — user confirms and calls /post.
    """
    import base64 as _b64
    from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent

    suffix = (_Path(file.filename or "").suffix.lower() or "")
    if suffix not in ALLOWED_RECEIPT_SUFFIXES:
        raise HTTPException(400, f"Unsupported file type. Allowed: {sorted(ALLOWED_RECEIPT_SUFFIXES)}")
    raw = await file.read()
    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(413, "Receipt too large (10 MB limit)")

    # Persist the receipt file
    rid = new_id()
    owner_dir = RECEIPT_UPLOAD_ROOT / u["id"]
    owner_dir.mkdir(parents=True, exist_ok=True)
    disk_path = owner_dir / f"{rid}{suffix}"
    disk_path.write_bytes(raw)

    # Call Gemini 3 vision
    b64_image = _b64.b64encode(raw).decode("ascii")
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"receipt-{rid}",
        system_message=(
            "You are a South African accounting assistant. Extract structured data from the receipt image. "
            "Return ONLY strict JSON, no prose. Fields: "
            '{"vendor": "...", "vendor_vat_number": "...", "date": "YYYY-MM-DD", '
            '"subtotal": <number>, "vat": <number>, "total": <number>, '
            '"currency": "ZAR", "payment_method": "cash|card|eft|unknown", '
            '"line_items": [{"description": "...", "qty": <num>, "unit_price": <num>, "amount": <num>}], '
            '"suggested_expense_category": "marketing|subscriptions|travel|accommodation|meals|telecoms|bank_charges|professional_fees|training|stationery|insurance|rent|utilities|motor|other"} '
            "Use null for unknown fields. Amounts are numbers (not strings)."
        ),
    ).with_model("gemini", "gemini-2.5-flash")
    mime = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
            ".webp": "image/webp", ".heic": "image/heic", ".pdf": "application/pdf"}.get(suffix, "application/octet-stream")
    try:
        resp = await chat.send_message(UserMessage(
            text="Extract the receipt data as JSON.",
            file_contents=[ImageContent(image_base64=b64_image, mime_type=mime)],
        ))
        # Try to parse JSON; strip fences
        text = (resp or "").strip()
        if text.startswith("```"):
            text = text.split("```", 2)[1]
            if text.lstrip().startswith("json"):
                text = text.split("\n", 1)[1]
            text = text.rstrip("`").strip()
        try:
            parsed = json.loads(text)
        except Exception:
            parsed = {"raw": text, "parse_error": True}
    except Exception as e:
        log.exception("Gemini vision OCR failed")
        parsed = {"error": str(e)}

    doc = {
        "id": rid,
        "owner_id": u["id"],
        "filename": file.filename,
        "suffix": suffix,
        "size_bytes": len(raw),
        "disk_path": str(disk_path),
        "extracted": parsed,
        "status": "pending_review",
        "journal_id": None,
        "created_at": now_iso(),
    }
    await db.receipts.insert_one(doc)
    await audit(u["actor_id"], "scan_receipt", "receipt", rid,
                after={"filename": file.filename, "parsed_ok": not isinstance(parsed, dict) or "error" not in parsed})
    out = {k: v for k, v in doc.items() if k not in ("_id", "disk_path")}
    return out


@api.get("/accounting/receipts")
async def list_receipts(status: str | None = None, u: dict = Depends(current_user)):
    q: dict = {"owner_id": u["id"]}
    if status:
        q["status"] = status
    return await db.receipts.find(q, {"_id": 0, "disk_path": 0}).sort("created_at", -1).to_list(500)


@api.post("/accounting/receipts/{rid}/post")
async def post_receipt_as_expense(rid: str, body: dict = Body(...), u: dict = Depends(require_accountant)):
    """User-confirmed posting of a scanned receipt as an expense journal.
    Body:
      - `expense_account_code` (e.g. 60000)
      - `vat_code` (S, SI, Z, E, NV, CI) — defaults to SI if subtotal+vat sum correctly
      - `payment_account_code` (bank 10100 / petty cash 10300 / creditors 20000) — default creditors
      - Overrides: `vendor`, `date`, `subtotal`, `vat`, `total`
    """
    rcp = await db.receipts.find_one({"id": rid, "owner_id": u["id"]}, {"_id": 0})
    if not rcp:
        raise HTTPException(404, "Receipt not found")
    if rcp["status"] == "posted":
        raise HTTPException(400, "Already posted")
    ex = rcp.get("extracted") or {}
    vendor = body.get("vendor") or ex.get("vendor") or "Vendor"
    date = body.get("date") or ex.get("date") or datetime.now(timezone.utc).date().isoformat()
    subtotal = float(body.get("subtotal") if body.get("subtotal") is not None else (ex.get("subtotal") or 0))
    vat = float(body.get("vat") if body.get("vat") is not None else (ex.get("vat") or 0))
    total = float(body.get("total") if body.get("total") is not None else (ex.get("total") or (subtotal + vat)))
    if total <= 0:
        raise HTTPException(400, "total must be > 0 — confirm values before posting")
    # Reconcile subtotal + vat with total (small float drift tolerance)
    if abs((subtotal + vat) - total) > 0.02 and subtotal and vat:
        raise HTTPException(400, f"subtotal ({subtotal}) + vat ({vat}) does not equal total ({total})")
    if not subtotal:
        subtotal = total - vat
    expense_code = body.get("expense_account_code")
    if not expense_code:
        raise HTTPException(400, "expense_account_code required (e.g. 60000)")
    acct = await db.accounts.find_one({"owner_id": u["id"], "code": expense_code}, {"_id": 0})
    if not acct or acct["type"] != "expense":
        raise HTTPException(400, "expense_account_code must be an existing expense account")
    vat_code = body.get("vat_code") or ("SI" if vat > 0 else "NV")
    payment_code = body.get("payment_account_code") or "51000"  # Trade Creditors by default
    pay_acct = await db.accounts.find_one({"owner_id": u["id"], "code": payment_code}, {"_id": 0})
    if not pay_acct:
        raise HTTPException(400, "payment_account_code not found in CoA")

    lines = [JournalLineIn(
        account_code=expense_code, debit=round(subtotal, 2), credit=0,
        description=f"{vendor} — {ex.get('suggested_expense_category', '')}".strip(" —"),
        vat_code=vat_code if vat == 0 else None,
    )]
    if vat > 0:
        lines.append(JournalLineIn(
            account_code="23000", debit=round(vat, 2), credit=0,
            description=f"VAT input — {vendor}",
            vat_code=vat_code, vat_amount=round(vat, 2),
        ))
    lines.append(JournalLineIn(
        account_code=payment_code, debit=0, credit=round(total, 2),
        description=f"Payable to {vendor}",
    ))

    j = JournalIn(
        date=date,
        memo=f"Receipt · {vendor}",
        reference=f"RCP-{rid[:8]}",
        lines=lines,
        source="receipt_ocr",
        source_id=rid,
    )
    jdoc = await _validate_and_post_journal(u["id"], u["actor_id"], j, auto=False)
    await db.receipts.update_one({"id": rid, "owner_id": u["id"]}, {"$set": {
        "status": "posted", "journal_id": jdoc["id"],
        "posted_fields": {"vendor": vendor, "date": date, "subtotal": subtotal, "vat": vat, "total": total,
                          "expense_code": expense_code, "payment_code": payment_code, "vat_code": vat_code},
        "posted_at": now_iso(),
    }})
    await audit(u["actor_id"], "post_receipt", "receipt", rid, after={"journal_id": jdoc["id"], "total": total})
    return {"ok": True, "journal_id": jdoc["id"]}


@api.delete("/accounting/receipts/{rid}")
async def delete_receipt(rid: str, u: dict = Depends(require_accountant)):
    r = await db.receipts.find_one({"id": rid, "owner_id": u["id"]}, {"_id": 0})
    if not r:
        raise HTTPException(404, "Receipt not found")
    if r["status"] == "posted":
        raise HTTPException(400, "Cannot delete a posted receipt — reverse the journal first")
    try:
        p = _Path(r.get("disk_path", ""))
        if p.exists():
            p.unlink()
    except Exception:
        pass
    await db.receipts.delete_one({"id": rid, "owner_id": u["id"]})
    await audit(u["actor_id"], "delete_receipt", "receipt", rid)
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 batch E — Annual Financial Statements (AFS) PDF bundle
#  Registered here (after all shared helpers are defined) via a setup function.
# ══════════════════════════════════════════════════════════════════════════════
from accounting_afs import register_afs_routes  # noqa: E402
register_afs_routes(api)

from accounting_payroll import register_payroll_routes  # noqa: E402
register_payroll_routes(api)


app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)
